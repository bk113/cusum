"""
CUSUM Active Manager Monitor — PDF Report Generator
====================================================
Page structure:
  1.    Cover
  2.    Executive Summary
  3.    Methodology
  4.    Model Assumptions & Parameters
  5-N.  Fund Universe Grid (6/page, 3×2)
  N+1.  Summary Results Table
  N+2.  Alarm Log
  N+3.  Framework Health Diagnostics
  N+4.  Waterfall Chart
  N+5.  Insufficient History Funds
  N+6.  Data Quality Appendix
  Last. Conclusions & Action Items

Run:
    python report.py
    python report.py --refresh-all
    python report.py --use-cache
    python report.py --output my.pdf
"""

import argparse
import io
import json
import math
import os
import pickle
import warnings
from datetime import datetime

import matplotlib
matplotlib.use('Agg')
import matplotlib.gridspec as gridspec
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.patches import FancyBboxPatch

from cusum import (
    ASSET_CLASS_PRESETS, CUSUMMonitor, MIN_MONITORING_MONTHS,
    PRESET_PAIRS, THRESHOLD_TABLE, TICKER_ASSET_CLASS, VIX_REGIME_DISPLAY,
    bootstrap_ir_ci, compute_extra_stats, load_from_yfinance,
    load_vix_history, monitor_fund, post_alarm_returns,
    sanity_check_returns, signal_quality, tag_alarm_regime,
)

# ── python-docx for Word export (optional; pip install python-docx) ───────────
try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ── Layout constants ──────────────────────────────────────────────────────────
FUNDS_PER_GRID_PAGE = 6
SUMMARY_PER_PAGE    = 18
ALARM_LOG_PER_PAGE  = 16
CONCLUSION_PER_PAGE = 14

# ── Colour palette (spec) ─────────────────────────────────────────────────────
NAVY   = "#1E2761"
CORAL  = "#E85D24"   # alarm / red
TEAL   = "#0F766E"   # healthy / green
AMBER  = "#BA7517"   # watchlist / elevated
GOLD   = "#D4AC0D"   # cover accent
GREY   = "#BDC3C7"
WHITE  = "#FFFFFF"
LGREY  = "#F4F6F7"
CARD   = "#F8FAFC"   # light card background
BORDER = "#E2E8F0"

# Legacy aliases kept for backward compat
RED   = CORAL
GREEN = TEAL

# ── Word document design constants ───────────────────────────────────────────
WORD_STYLE = {
    "navy_hex":        "1E2761",
    "coral_hex":       "E85D24",
    "teal_hex":        "0F766E",
    "amber_hex":       "BA7517",
    "gold_hex":        "D4AC0D",
    "grey_hex":        "BDC3C7",
    "lgrey_hex":       "F4F6F7",
    "header_font":     "Calibri",
    "body_font":       "Calibri",
    "heading_font":    "Calibri",
    "cover_title_pt":  28,
    "heading1_pt":     16,
    "heading2_pt":     13,
    "body_pt":         10,
    "table_pt":        9,
    "page_header_tpl": "CUSUM Active Manager Monitor  |  Data as of {date}",
    "confidential":    "CONFIDENTIAL — For internal use only",
}

# ── Fund metadata for Word catalogue ─────────────────────────────────────────
# (inception_year, aum_band, typical_te_range)
# AUM bands: Large >$10B, Mid $1–10B, Small <$1B.
# Source: fund prospectuses / publicly available data as of template creation.
FUND_METADATA: dict = {
    # ── AIAIM Singapore-listed ────────────────────────────────────────────────
    "0P00008T6D.SI": (2015, "Small", "3-5%"),
    "0P00008T6G.SI": (2015, "Small", "3-5%"),
    "0P00008T6N.SI": (2015, "Small", "3-5%"),
    "0P0001I6KI.SI": (2020, "Small", "3-5%"),
    "0P0001KL95.SI": (2019, "Small", "3-5%"),
    "0P0001RMIN.SI": (2021, "Small", "3-5%"),
    # ── Balanced ─────────────────────────────────────────────────────────────
    "PRWCX": (1986, "Large", "3-5%"),
    "VWELX": (1929, "Large", "2-4%"),
    # ── Developed International ───────────────────────────────────────────────
    "DODWX": (2001, "Large", "4-7%"),
    "OAKIX": (1992, "Mid",   "5-8%"),
    # ── EM Hard Currency Bond ─────────────────────────────────────────────────
    "DBLEX": (2014, "Small", "3-5%"),
    "EBNDX": (2007, "Mid",   "3-5%"),
    "IEMIX": (2002, "Mid",   "3-5%"),
    "MSEDX": (1998, "Small", "3-5%"),
    "PEBIX": (1997, "Large", "3-5%"),
    "PREMX": (1994, "Mid",   "3-5%"),
    "TGBAX": (1986, "Mid",   "4-6%"),
    "VWOB":  (2013, "Mid",   "3-4%"),
    # ── EM Local Currency Bond ────────────────────────────────────────────────
    "PFMIX": (1992, "Mid",   "6-10%"),
    "PRLAX": (2008, "Small", "7-11%"),
    # ── EM Equity ─────────────────────────────────────────────────────────────
    "ABEMX": (1992, "Mid",   "5-9%"),
    "BGEFX": (2010, "Small", "6-10%"),
    "EEMAX": (2008, "Small", "6-9%"),
    "FSEAX": (1993, "Mid",   "5-9%"),
    "GERIX": (1997, "Mid",   "5-9%"),
    "HAEMX": (1994, "Small", "5-9%"),
    "HLEMX": (2009, "Mid",   "5-8%"),
    "HLMEX": (1998, "Mid",   "5-8%"),
    "JEMWX": (1992, "Mid",   "5-9%"),
    "LZOEX": (1994, "Mid",   "5-9%"),
    "MADCX": (2010, "Small", "5-9%"),
    "MEMCX": (2007, "Small", "6-10%"),
    "MIEMX": (1995, "Mid",   "5-9%"),
    "NEWFX": (2000, "Mid",   "5-8%"),
    "NWFFX": (1999, "Large", "5-8%"),
    "ODMAX": (1994, "Mid",   "5-9%"),
    "PDEZX": (2007, "Small", "5-9%"),
    "PEIFX": (1999, "Small", "6-10%"),
    "PRMSX": (1995, "Mid",   "5-9%"),
    "REMIX": (2010, "Small", "6-10%"),
    "SFENX": (2012, "Small", "5-9%"),
    "TEDMX": (1991, "Mid",   "5-9%"),
    "TEMFX": (1988, "Small", "5-9%"),
    "VEIEX": (1994, "Large", "5-8%"),
    "VEMAX": (2001, "Large", "5-8%"),
    "WAEMX": (2006, "Small", "6-9%"),
    "WIEMX": (2005, "Small", "6-9%"),
    # ── Global Equity (ACWI) ──────────────────────────────────────────────────
    "CWGIX": (1993, "Large", "3-5%"),
    "GPROX": (2006, "Small", "4-6%"),
    "MGGIX": (1986, "Mid",   "3-5%"),
    "OPGIX": (1969, "Mid",   "4-6%"),
    "SGIGX": (2003, "Small", "4-6%"),
    "TEDIX": (1992, "Mid",   "4-6%"),
    "VHGEX": (1995, "Mid",   "3-5%"),
    # ── Global Equity (URTH / MSCI World) ────────────────────────────────────
    "AIVSX": (1934, "Mid",   "3-5%"),
    "HAINX": (1987, "Mid",   "4-6%"),
    "JGVAX": (1998, "Small", "4-6%"),
    "MGIAX": (1996, "Mid",   "4-6%"),
    # ── US Fixed Income ───────────────────────────────────────────────────────
    "LSBRX": (1991, "Large", "1-2%"),
    "PTTRX": (1987, "Large", "0.5-1.5%"),
    # ── US Large-Cap ─────────────────────────────────────────────────────────
    "AGTHX": (1973, "Large", "4-6%"),
    "DODGX": (1965, "Large", "4-6%"),
    "FCNTX": (1967, "Large", "4-6%"),
    "FMAGX": (1963, "Mid",   "4-6%"),
    "SEQUX": (1970, "Mid",   "5-7%"),
    # ── US Small-Cap ─────────────────────────────────────────────────────────
    "FLPSX": (1989, "Large", "5-8%"),
    "OTCFX": (1956, "Mid",   "5-8%"),
}

# ── Asset class lookup ────────────────────────────────────────────────────────
BENCH_TO_ASSET_CLASS = {
    "IWF":  "US Large Growth",
    "IWD":  "US Large Value",
    "SPY":  "US Large Blend",
    "IWM":  "US Small Cap",
    "EFA":  "Intl Developed",
    "EEM":  "Emerging Markets",
    "AOR":  "Balanced 60/40",
    "AGG":  "US Fixed Income",
    "EMB":  "EM Hard CCY Debt",
    "EMLC": "EM Local CCY Debt",
    "ACWI": "Global Equity",
    "URTH": "Global Equity",
}
TICKER_AC_OVERRIDES = {
    "FLPSX": "US Small/Mid Cap",
    "DODWX": "Intl Value",
}
ASSET_CLASS_ORDER = [
    "US Large Blend", "US Large Growth", "US Large Value",
    "US Small Cap", "US Small/Mid Cap",
    "Balanced 60/40", "US Fixed Income",
    "Intl Developed", "Intl Value",
    "Emerging Markets", "EM Hard CCY Debt", "EM Local CCY Debt",
    "Global Equity",
]
ASSET_CLASS_COLORS = {
    "US Large Blend":    NAVY,
    "US Large Growth":   NAVY,
    "US Large Value":    NAVY,
    "US Small Cap":      NAVY,
    "US Small/Mid Cap":  NAVY,
    "Intl Developed":    "#1A5276",
    "Intl Value":        "#1A5276",
    "Emerging Markets":  "#1A5276",
    "Balanced 60/40":    "#145A32",
    "US Fixed Income":   "#145A32",
    "EM Hard CCY Debt":  "#6E2F8A",
    "EM Local CCY Debt": "#6E2F8A",
    "Global Equity":     "#7D6608",
}
AC_TO_PRESET = {
    "US Large Blend":    "equity",
    "US Large Growth":   "equity",
    "US Large Value":    "equity",
    "US Small Cap":      "equity",
    "US Small/Mid Cap":  "equity",
    "Intl Developed":    "equity",
    "Intl Value":        "equity",
    "Emerging Markets":  "equity",
    "Global Equity":     "equity",
    "Balanced 60/40":    "balanced",
    "US Fixed Income":   "fixed_income",
    "EM Hard CCY Debt":  "em_debt",
    "EM Local CCY Debt": "em_debt",
}
# Governance h per asset class (from spec)
GOVERNANCE_H = {
    "equity":       19.46,
    "balanced":     20.26,
    "fixed_income": 21.16,
    "em_debt":      21.16,
}

CALIB_MONTHS = 24


# ── Pure helpers ──────────────────────────────────────────────────────────────

def _asset_class(fund_ticker, bench):
    return TICKER_AC_OVERRIDES.get(fund_ticker,
           BENCH_TO_ASSET_CLASS.get(bench, "Other"))


def _ac_sort_key(ac):
    try:
        return ASSET_CLASS_ORDER.index(ac)
    except ValueError:
        return len(ASSET_CLASS_ORDER)


def _preset_for(fund_ticker, bench):
    ac = _asset_class(fund_ticker, bench)
    return AC_TO_PRESET.get(ac, "equity")


def _h_for(fund_ticker, bench):
    preset = _preset_for(fund_ticker, bench)
    return ASSET_CLASS_PRESETS.get(preset, {}).get("threshold",
           GOVERNANCE_H.get(preset, 19.46))


def _fmt_pct(v, decimals=1):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "n/a"
    return f"{v:+.{decimals}%}"


def _fmt_num(v, decimals=2):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "n/a"
    return f"{v:+.{decimals}f}"


# ── Pagination helpers ────────────────────────────────────────────────────────

def _paginate_grouped_rows(rows, per_page):
    pages, current, used = [], [], 0
    for kind, data in rows:
        if used >= per_page:
            pages.append(current); current = []; used = 0
        if kind == "header" and used == per_page - 1:
            pages.append(current); current = []; used = 0
        current.append((kind, data)); used += 1
    if current:
        pages.append(current)
    return pages or [[]]


def _build_grouped_rows_summary(all_results):
    sorted_res = sorted(all_results,
                        key=lambda r: (_ac_sort_key(_asset_class(r["fund"], r["bench"])),
                                       r["fund"]))
    rows, prev_ac = [], None
    for r in sorted_res:
        ac = _asset_class(r["fund"], r["bench"])
        if ac != prev_ac:
            rows.append(("header", ac))
            prev_ac = ac
        rows.append(("data", r))
    return rows


# ── Page count ────────────────────────────────────────────────────────────────

def compute_total_pages(all_results, insufficient_history):
    n_loaded = len(all_results)
    n_alarm  = sum(1 for r in all_results if r["alarm"])
    n_insuff = len(insufficient_history)

    # count all alarms across all funds for alarm log
    total_alarm_events = sum(r.get("n_alarms", 0) for r in all_results)

    summ_rows  = _build_grouped_rows_summary(all_results)
    n_summ     = max(1, len(_paginate_grouped_rows(summ_rows, SUMMARY_PER_PAGE)))
    n_grid     = max(1, math.ceil(n_loaded / FUNDS_PER_GRID_PAGE))
    n_alarm_log= max(1, math.ceil(max(total_alarm_events, 1) / ALARM_LOG_PER_PAGE))
    n_insuff_p = max(1, math.ceil(max(n_insuff, 1) / 20)) if n_insuff > 0 else 1
    dq_funds   = [r for r in all_results if r.get("dq_flags")]
    n_dq       = max(1, math.ceil(max(len(dq_funds), 1) / 20))

    return (
        1            # cover
        + 1          # exec summary
        + 1          # methodology
        + 1          # assumptions
        + n_grid     # fund universe grid
        + n_summ     # summary results table
        + n_alarm_log# alarm log
        + 1          # framework health
        + 1          # waterfall
        + n_insuff_p # insufficient history
        + n_dq       # data quality
        + 1          # conclusions
    )


# ── Footer ────────────────────────────────────────────────────────────────────

def page_footer(fig, page_num, total_pages):
    today = datetime.today().strftime("%d %b %Y")
    fig.text(0.5, 0.018,
             f"CUSUM Active Manager Monitor  |  Page {page_num} of {total_pages}  |  Generated {today}",
             ha="center", va="bottom", fontsize=7.5, color="grey")
    fig.text(0.07, 0.018,
             "CONFIDENTIAL — For internal use only",
             ha="left", va="bottom", fontsize=7, color="grey", style="italic")


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — Cover
# ═══════════════════════════════════════════════════════════════════════════════

def page_cover(pdf, pnum, total_pages):
    today     = datetime.today()
    month_str = today.strftime("%B %Y")
    cutoff    = (today.replace(day=1) - pd.Timedelta(days=1)).strftime("%d %B %Y")

    fig = plt.figure(figsize=(11, 8.5))
    fig.patch.set_facecolor(NAVY)

    fig.text(0.5, 0.88, "CUSUM Active Manager Monitor",
             ha="center", fontsize=28, fontweight="bold", color=WHITE)
    fig.text(0.5, 0.815, f"Monthly Governance Report — {month_str}",
             ha="center", fontsize=15, color=GOLD)

    fig.add_artist(plt.Line2D([0.08, 0.92], [0.79, 0.79],
                              transform=fig.transFigure, color=GOLD, lw=1.2))

    fig.text(0.5, 0.755, f"Report Date: {today.strftime('%d %B %Y')}   |   Data Cutoff: {cutoff}",
             ha="center", fontsize=10, color=GREY)
    fig.text(0.5, 0.725, "Universe: 67 active mandates across 9 asset classes",
             ha="center", fontsize=10, color=WHITE)

    # Threshold table
    fig.text(0.5, 0.685, "Governance Alarm Thresholds (h) by Asset Class",
             ha="center", fontsize=11, fontweight="bold", color=GOLD)

    th_data = [
        ("Equity (all mandates)",          "19.46"),
        ("Balanced (60/40)",               "20.26"),
        ("Fixed Income (US bonds)",        "21.16"),
        ("EM Debt (hard & local CCY)",     "21.16"),
    ]
    col_xs = [0.25, 0.62]
    row_y  = 0.650
    for i, (label, h_val) in enumerate(th_data):
        bg = "#2A3580" if i % 2 == 0 else "#242D6E"
        fig.add_artist(FancyBboxPatch((0.18, row_y - 0.025), 0.64, 0.026,
                                      boxstyle="round,pad=0.002",
                                      facecolor=bg, edgecolor="none",
                                      transform=fig.transFigure))
        fig.text(col_xs[0], row_y - 0.012, label, ha="center", va="center",
                 fontsize=9, color=WHITE, transform=fig.transFigure)
        fig.text(col_xs[1], row_y - 0.012, f"h = {h_val}", ha="center", va="center",
                 fontsize=9, color=GOLD, fontweight="bold", transform=fig.transFigure)
        row_y -= 0.028

    fig.add_artist(plt.Line2D([0.08, 0.92], [0.50, 0.50],
                              transform=fig.transFigure, color=GOLD, lw=0.8))

    fig.text(0.08, 0.470, "What is CUSUM?", fontsize=12, fontweight="bold", color=GOLD)
    fig.text(0.08, 0.430,
             "CUSUM (Cumulative Sum) is a sequential change-point detection method from statistical\n"
             "process control, adapted for investment management by Philips, Yashchin & Stein (2003).\n"
             "It detects persistent skill loss earlier and more rigorously than trailing 3-year IR reviews.",
             fontsize=9, color=WHITE, va="top", linespacing=1.6)

    bullets = [
        ("ARL Control",      "h calibrated by Monte Carlo to target ARL = 60 monitoring months (1 false alarm / 5 yrs per fund)."),
        ("VIX Regime Tags",  "Each alarm labelled calm / elevated / crisis based on prevailing VIX at alarm date."),
        ("Rolling Restart",  "L resets to 0 after each alarm — monitoring never stops; all events logged."),
        ("Actionable",       "An alarm triggers a structured IC review, NOT automatic termination."),
    ]
    by = 0.360
    for title, body in bullets:
        fig.text(0.10, by, f"\u2022  {title}:", fontsize=8.5, fontweight="bold", color=GOLD)
        fig.text(0.27, by, body, fontsize=8.2, color=GREY, va="top", linespacing=1.35)
        by -= 0.048

    fig.text(0.5, 0.140, "Prepared by: Investment Risk & Analytics",
             ha="center", fontsize=10, color=WHITE, fontstyle="italic")

    fig.text(0.5, 0.045,
             "Reference: Philips, T., Yashchin, E. & Stein, D. (2003). "
             "\"Using Statistical Process Control to Monitor Active Managers.\" "
             "Journal of Portfolio Management, 30(1), 86–94.",
             ha="center", fontsize=7.5, color=GREY, style="italic")

    page_footer(fig, pnum[0], total_pages)
    pnum[0] += 1
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — Executive Summary
# ═══════════════════════════════════════════════════════════════════════════════

def page_executive_summary(pdf, all_results, all_post_alarm, vix_series, pnum, total_pages):
    today    = datetime.today()
    jan1     = pd.Timestamp(today.year, 1, 1)
    cur_mo_s = pd.Timestamp(today.year, today.month, 1)
    cur_mo_e = (cur_mo_s + pd.offsets.MonthEnd(0))

    fig = plt.figure(figsize=(11, 8.5))
    fig.patch.set_facecolor(WHITE)

    fig.text(0.5, 0.960, "Executive Summary", ha="center",
             fontsize=20, fontweight="bold", color=NAVY)
    fig.text(0.5, 0.932, f"CUSUM Active Manager Monitor  •  {today.strftime('%d %B %Y')}",
             ha="center", fontsize=11, color="grey")

    # ── Stats ─────────────────────────────────────────────────────────────────
    n_total = len(all_results)

    # funds currently in alarm (L_now > h)
    in_alarm_now = [r for r in all_results
                    if r.get("L_now", 0) > r.get("h_threshold", 19.46)]
    # watchlist (L_now > 0.75*h)
    watchlist = [r for r in all_results
                 if r.get("L_now", 0) > 0.75 * r.get("h_threshold", 19.46)
                 and r not in in_alarm_now]

    # new alarms this month
    def _alarm_in_month(r):
        ts = r.get("alarm_date_ts")
        if ts is None:
            return False
        return cur_mo_s <= ts <= cur_mo_e

    new_alarms_this_month = [r for r in all_results if _alarm_in_month(r)]

    # YTD alarms
    def _alarm_ytd(r):
        ts = r.get("alarm_date_ts")
        if ts is None:
            return False
        return ts >= jan1

    ytd_alarms = [r for r in all_results if _alarm_ytd(r)]

    # Regime breakdown for new alarms
    def _regime_of(r):
        return r.get("vix_regime", "unknown")

    regime_counts_new = {}
    for r in new_alarms_this_month:
        reg = _regime_of(r)
        regime_counts_new[reg] = regime_counts_new.get(reg, 0) + 1

    # YTD regime breakdown
    ytd_regime = {}
    for r in ytd_alarms:
        reg = _regime_of(r)
        ytd_regime[reg] = ytd_regime.get(reg, 0) + 1

    # framework health: calm-regime signal quality
    calm_sq = None
    sq_vals = []
    for r in all_results:
        sq = r.get("sig_quality")
        if sq is not None and r.get("vix_regime") == "calm":
            sq_vals.append(sq)
    if sq_vals:
        calm_sq = float(np.mean(sq_vals))

    # KPI boxes
    kpi_data = [
        (str(len(new_alarms_this_month)), "New Alarms\nThis Month",
         CORAL if new_alarms_this_month else TEAL),
        (str(len(in_alarm_now)),          "Currently\nIn Alarm",
         CORAL if in_alarm_now else TEAL),
        (str(len(watchlist)),             "Watchlist\n(L > 75% h)",
         AMBER if watchlist else TEAL),
        (str(len(ytd_alarms)),            "YTD Alarms\n(since Jan 1)",
         CORAL if ytd_alarms else TEAL),
    ]
    box_xs  = [0.030, 0.270, 0.510, 0.750]
    box_w   = 0.210
    box_top = 0.895
    box_bot = 0.800

    for (val, label, col), bx in zip(kpi_data, box_xs):
        fig.add_artist(FancyBboxPatch((bx, box_bot), box_w, box_top - box_bot,
                                      boxstyle="round,pad=0.008",
                                      facecolor=col, edgecolor="none",
                                      transform=fig.transFigure, alpha=0.10))
        fig.add_artist(FancyBboxPatch((bx, box_bot), box_w, box_top - box_bot,
                                      boxstyle="round,pad=0.008",
                                      facecolor="none", edgecolor=col, lw=2,
                                      transform=fig.transFigure))
        mid = bx + box_w / 2
        fig.text(mid, box_bot + (box_top - box_bot) * 0.62, val,
                 ha="center", va="center", fontsize=26, fontweight="bold", color=col)
        fig.text(mid, box_bot + (box_top - box_bot) * 0.22, label,
                 ha="center", va="center", fontsize=8.5, color="#555555", linespacing=1.4)

    y_cur = 0.785

    # ── New alarms this month breakdown ──────────────────────────────────────
    fig.text(0.030, y_cur, f"New Alarms This Month — Regime Breakdown",
             fontsize=9.5, fontweight="bold", color=NAVY)
    y_cur -= 0.026
    if new_alarms_this_month:
        regime_str = "  |  ".join(
            f"{reg.capitalize()}: {cnt}" for reg, cnt in sorted(regime_counts_new.items()))
        funds_str  = ", ".join(r["fund"] for r in new_alarms_this_month[:8])
        if len(new_alarms_this_month) > 8:
            funds_str += " ..."
        fig.text(0.040, y_cur, f"Regimes: {regime_str}     Funds: {funds_str}",
                 fontsize=8.2, color="#333333")
    else:
        fig.text(0.040, y_cur, "No new alarms this month.", fontsize=8.2, color=TEAL)
    y_cur -= 0.030

    # ── Currently in alarm ────────────────────────────────────────────────────
    fig.text(0.030, y_cur, "Funds Currently in Alarm State  (L\u209c \u2265 h)",
             fontsize=9.5, fontweight="bold", color=CORAL)
    y_cur -= 0.024
    if in_alarm_now:
        for r in in_alarm_now[:6]:
            h = r.get("h_threshold", 19.46)
            L = r.get("L_now", 0.0)
            fig.text(0.042, y_cur,
                     f"\u2022  {r['fund']} ({r['bench']})  —  L={L:.2f}  h={h:.2f}  "
                     f"alarm: {r.get('alarm_date','-')}  regime: {r.get('vix_regime','-')}",
                     fontsize=7.8, color=CORAL)
            y_cur -= 0.020
        if len(in_alarm_now) > 6:
            fig.text(0.042, y_cur,
                     f"  ... and {len(in_alarm_now)-6} more (see Summary Table)",
                     fontsize=7.5, color="#777777")
            y_cur -= 0.020
    else:
        fig.text(0.042, y_cur, "None — all funds below alarm threshold.",
                 fontsize=8, color=TEAL)
        y_cur -= 0.020
    y_cur -= 0.012

    # ── Watchlist ─────────────────────────────────────────────────────────────
    fig.text(0.030, y_cur, "Watchlist  (L\u209c \u2265 0.75 \u00d7 h)",
             fontsize=9.5, fontweight="bold", color=AMBER)
    y_cur -= 0.024
    if watchlist:
        wl_str = "  |  ".join(
            f"{r['fund']} (L={r.get('L_now',0):.1f})" for r in watchlist[:8])
        if len(watchlist) > 8:
            wl_str += f" ... +{len(watchlist)-8} more"
        fig.text(0.042, y_cur, wl_str, fontsize=7.8, color=AMBER)
    else:
        fig.text(0.042, y_cur, "No funds approaching alarm threshold.", fontsize=8, color=TEAL)
    y_cur -= 0.030

    # ── YTD alarm activity ────────────────────────────────────────────────────
    fig.text(0.030, y_cur, f"YTD Alarm Activity ({today.year})",
             fontsize=9.5, fontweight="bold", color=NAVY)
    y_cur -= 0.024
    ytd_str = f"Total: {len(ytd_alarms)}"
    if ytd_regime:
        ytd_str += "  |  " + "  |  ".join(
            f"{r.capitalize()}: {c}" for r, c in sorted(ytd_regime.items()))
    fig.text(0.042, y_cur, ytd_str, fontsize=8.2, color="#333333")
    y_cur -= 0.030

    # ── Framework health ──────────────────────────────────────────────────────
    fig.text(0.030, y_cur, "Framework Health  —  Calm-Regime Signal Quality",
             fontsize=9.5, fontweight="bold", color=NAVY)
    y_cur -= 0.024
    if calm_sq is not None:
        sq_col = TEAL if calm_sq >= 0.60 else (CORAL if calm_sq < 0.55 else AMBER)
        sq_flag = ""
        if calm_sq < 0.55:
            sq_flag = "  ⚠ BELOW 55% — CONSIDER RECALIBRATION"
        fig.text(0.042, y_cur,
                 f"Calm-regime signal quality (% correct at +12M): {calm_sq:.0%}{sq_flag}",
                 fontsize=8.2, color=sq_col, fontweight="bold" if sq_flag else "normal")
    else:
        fig.text(0.042, y_cur, "Signal quality pending (insufficient post-alarm data).",
                 fontsize=8.2, color="#777777")
    y_cur -= 0.030

    # ── Recommended actions ───────────────────────────────────────────────────
    fig.text(0.030, y_cur, "Recommended Actions — Funds Requiring CIO Attention",
             fontsize=9.5, fontweight="bold", color=NAVY)
    y_cur -= 0.024
    action_funds = in_alarm_now + [r for r in watchlist if r not in in_alarm_now]
    if action_funds:
        for r in action_funds[:8]:
            state = "IN ALARM" if r in in_alarm_now else "WATCHLIST"
            col   = CORAL if r in in_alarm_now else AMBER
            fig.text(0.042, y_cur,
                     f"\u2022  [{state}]  {r['fund']} — {r.get('name', '')}",
                     fontsize=7.8, color=col)
            y_cur -= 0.020
    else:
        fig.text(0.042, y_cur, "No funds currently require CIO attention.",
                 fontsize=8, color=TEAL)

    page_footer(fig, pnum[0], total_pages)
    pnum[0] += 1
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — Methodology
# ═══════════════════════════════════════════════════════════════════════════════

def page_methodology(pdf, pnum, total_pages):
    fig = plt.figure(figsize=(11, 8.5))
    fig.patch.set_facecolor(WHITE)
    fig.text(0.5, 0.962, "Methodology — How CUSUM Works", ha="center",
             fontsize=20, fontweight="bold", color=NAVY)
    fig.text(0.5, 0.932,
             "Philips-Yashchin-Stein (2003) sequential change-point detection for active managers",
             ha="center", fontsize=10, color="grey")

    # Four building blocks
    fig.text(0.040, 0.900, "The Four Building Blocks", fontsize=11, fontweight="bold", color=NAVY)

    blocks = [
        ("1. Information Ratio (IR)",
         "IR_hat_t = (e_t / \u03c3_monthly) \u00d7 \u221a12",
         "Monthly log excess return standardised by EWMA tracking error, annualised. "
         "Std \u2248 \u221a12 under H\u2080 — matches PYS (2003) Table 2 scale."),
        ("2. Par Level \u03bc = 0.25",
         "drift = 0.5 \u00d7 (\u03bc_good + \u03bc_bad) \u2212 IR_hat",
         "Drift parameter \u03bc = 0.25 (midpoint of good=0.5 and bad=0.0) is the "
         "Moustakides-optimal reference. Positive drift accumulates when IR_hat < 0.25."),
        ("3. Tally Statistic L",
         "L_t = max(0,  L_{t-1} + 0.25 \u2212 IR_hat_t)",
         "Lindley recursion. Floor at zero: L cannot go negative. "
         "L grows during underperformance, resets during outperformance. "
         "Each run of bad months compounds the evidence."),
        ("4. Threshold h (governance h)",
         "ALARM if  L_t \u2265 h;  then reset  L_t = 0",
         "h is calibrated by Monte Carlo (ARL target = 60 monitoring months = 5 years). "
         "Per-asset-class values: equity 19.46 | balanced 20.26 | bonds/EM-debt 21.16. "
         "After alarm, L resets and monitoring continues — rolling restart."),
    ]

    y = 0.875
    for i, (title, formula, explanation) in enumerate(blocks):
        col = NAVY if i % 2 == 0 else "#1A5276"
        fig.add_artist(FancyBboxPatch((0.040, y - 0.068), 0.92, 0.072,
                                      boxstyle="round,pad=0.006",
                                      facecolor=CARD, edgecolor=BORDER, lw=0.8,
                                      transform=fig.transFigure))
        fig.text(0.052, y - 0.012, title, fontsize=9.5, fontweight="bold", color=col)
        fig.text(0.052, y - 0.030, formula, fontsize=8.5, color=CORAL,
                 fontfamily="monospace")
        fig.text(0.052, y - 0.048, explanation, fontsize=7.8, color="#444444",
                 va="top", linespacing=1.4)
        y -= 0.082

    # VIX regime table
    y -= 0.010
    fig.text(0.040, y, "VIX Regime Classification", fontsize=10, fontweight="bold", color=NAVY)
    y -= 0.028

    vix_rows = [
        ("calm",     "VIX < 20",       TEAL,   "Normal market conditions. CUSUM signals most reliable."),
        ("elevated", "20 \u2264 VIX < 30", AMBER,  "Moderate stress. Some alarms may be regime-driven."),
        ("crisis",   "VIX \u2265 30",   CORAL,  "Tail-risk episode (e.g. GFC, COVID). Alarm may reflect market not skill."),
    ]
    for regime, cutoff, col, note in vix_rows:
        fig.add_artist(FancyBboxPatch((0.040, y - 0.022), 0.92, 0.024,
                                      boxstyle="round,pad=0.003",
                                      facecolor=col, edgecolor="none",
                                      transform=fig.transFigure, alpha=0.08))
        fig.text(0.055, y - 0.010, regime.upper(), fontsize=8.5, fontweight="bold", color=col)
        fig.text(0.160, y - 0.010, cutoff, fontsize=8.5, color="#333333")
        fig.text(0.310, y - 0.010, note, fontsize=8.0, color="#444444")
        y -= 0.028

    # Reference
    fig.text(0.5, 0.040,
             "Reference: Philips, T., Yashchin, E. & Stein, D. (2003). "
             "\"Using Statistical Process Control to Monitor Active Managers.\" "
             "Journal of Portfolio Management, 30(1), 86\u201394.",
             ha="center", fontsize=8, color="grey", style="italic")

    page_footer(fig, pnum[0], total_pages)
    pnum[0] += 1
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — Model Assumptions & Parameters (kept from original)
# ═══════════════════════════════════════════════════════════════════════════════

def page_assumptions(pdf, pnum, total_pages):
    fig = plt.figure(figsize=(11, 8.5))
    fig.patch.set_facecolor(WHITE)
    fig.text(0.5, 0.962, "Model Assumptions & Parameters", ha="center",
             fontsize=20, fontweight="bold", color=NAVY)
    fig.text(0.5, 0.932, "Philips-Yashchin-Stein (2003) — implementation reference",
             ha="center", fontsize=11, color="grey")

    fig.add_artist(plt.Line2D([0.508, 0.508], [0.055, 0.905],
                              transform=fig.transFigure, color=LGREY, lw=1.5))

    fig.text(0.040, 0.898, "Algorithm Parameters", fontsize=11, fontweight="bold", color=NAVY)

    params = [
        ("Threshold (h)",       "per-AC",      "equity 19.46 | balanced 20.26 | bonds/EM-debt 21.16"),
        ("EWMA decay (\u03b3)", "0.90",         "Von Neumann estimator; ~10-month half-life"),
        ("\u03bc_good (IR floor)", "+0.50",     "Acceptable manager skill (equity); 0.40/0.30 for bal/FI"),
        ("\u03bc_bad  (IR floor)", " 0.00",     "Underperforming manager"),
        ("Calibration window",  "24 months",    "Sigma warmup only; L frozen at zero"),
        ("Regime threshold",    "\u00b114.0",   "|IR_hat|>14: freeze L (~3.9\u03c3 with ann. IR std\u2248\u221a12)"),
        ("Bootstrap blocks",    "n^(1/3)",      "Circular block (Politis-White 2004)"),
        ("Min. monitoring",     "60 months",    "Below this: insufficient-history list"),
        ("ARL target",          "60 months",    "~1 false alarm per 5 years per fund"),
    ]

    p_col_x = [0.040, 0.185, 0.265]
    hdr_y   = 0.866
    fig.add_artist(FancyBboxPatch((0.035, hdr_y - 0.007), 0.458, 0.021,
                                   boxstyle="round,pad=0.002", facecolor=NAVY,
                                   edgecolor="none", transform=fig.transFigure))
    for hdr, x in zip(["Parameter", "Value", "Rationale"], p_col_x):
        fig.text(x, hdr_y, hdr, fontsize=8, fontweight="bold", color=WHITE, va="center")

    ROW_P = 0.050
    py    = hdr_y - 0.012
    for i, (param, val, rationale) in enumerate(params):
        py -= ROW_P
        bg = LGREY if i % 2 == 0 else WHITE
        fig.add_artist(FancyBboxPatch((0.035, py - 0.006), 0.458, ROW_P - 0.004,
                                      boxstyle="round,pad=0.002", facecolor=bg,
                                      edgecolor="none", transform=fig.transFigure))
        fig.text(p_col_x[0], py + ROW_P * 0.38, param,     fontsize=7.5, va="center", color=NAVY)
        fig.text(p_col_x[1], py + ROW_P * 0.38, val,       fontsize=7.5, va="center",
                 color=CORAL, fontweight="bold", fontfamily="monospace")
        fig.text(p_col_x[2], py + ROW_P * 0.20, rationale, fontsize=7.0, va="center",
                 color="#555555", linespacing=1.3)

    cls_y = py - 0.040
    fig.text(0.040, cls_y, "Alarm State Classification", fontsize=10, fontweight="bold", color=NAVY)
    cls_items = [
        (CORAL,  "IN ALARM",   "L\u209c \u2265 h. IC review required. Mandate compliance check."),
        (AMBER,  "WATCHLIST",  "L\u209c \u2265 0.75\u00d7h. Elevated monitoring. Revisit next month."),
        (TEAL,   "CLEAN",      "L\u209c < 0.75\u00d7h. No action required. Standard quarterly cadence."),
    ]
    cy = cls_y - 0.030
    for col, cls, desc in cls_items:
        fig.add_artist(FancyBboxPatch((0.038, cy - 0.005), 0.015, 0.016,
                                      boxstyle="round,pad=0.001", facecolor=col,
                                      edgecolor="none", transform=fig.transFigure))
        fig.text(0.060, cy + 0.003, f"{cls}:", fontsize=8, fontweight="bold", color=col, va="center")
        fig.text(0.150, cy + 0.003, desc, fontsize=7.5, color="#333333", va="center")
        cy -= 0.034

    # Right column
    fig.text(0.522, 0.898, "Data Source & Coverage", fontsize=11, fontweight="bold", color=NAVY)
    data_items = [
        ("Source",         "Yahoo Finance (yfinance) — no API key required"),
        ("Prices",         "Total-return adjusted close (dividends reinvested)"),
        ("Frequency",      "Daily prices \u2192 month-end to month-end returns"),
        ("Universe start", "January 2005 (common start for all funds)"),
        ("Returns used",   "Log excess: e\u209c = ln[(1+r_fund) / (1+r_bench)]"),
        ("VIX data",       "^VIX from yfinance, monthly MEAN of daily closes"),
        ("Benchmarks",     "SPY \u00b7 IWF \u00b7 IWD \u00b7 IWM \u00b7 EEM \u00b7 AGG \u00b7 EMB \u00b7 EMLC \u00b7 ACWI \u00b7 AOR"),
    ]
    ry = 0.856
    for label, val in data_items:
        fig.text(0.522, ry, f"{label}:", fontsize=8.5, color=NAVY, fontweight="bold", va="top")
        fig.text(0.648, ry, val, fontsize=8.2, color="#333333", va="top", linespacing=1.3)
        ry -= 0.037

    lim_y = ry - 0.015
    fig.text(0.522, lim_y, "Known Limitations", fontsize=11, fontweight="bold", color=NAVY)
    limitations = [
        "Multiple testing: 67 funds \u00d7 ARL=60m \u2192 ~13 expected false alarms/year across universe.",
        "Yahoo Finance adjusted prices may differ from official NAV series.",
        "Singapore AIAIM funds use Morningstar proxy IDs (0P...) with limited history.",
        "Regime outlier months (|IR_hat|>14): L frozen, not clipped \u2014 no distortion.",
        "Calibration sigma uses first 24m of actual data (mild look-ahead).",
    ]
    ly = lim_y - 0.028
    for i, lim in enumerate(limitations, 1):
        fig.text(0.522, ly, f"  {i}.  {lim}", fontsize=7.8, color="#333333",
                 va="top", linespacing=1.35)
        ly -= 0.035

    page_footer(fig, pnum[0], total_pages)
    pnum[0] += 1
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGES 5-N — Fund Universe Grid (6 per page, 3×2)
# ═══════════════════════════════════════════════════════════════════════════════

def _draw_grid_cell(fig, outer_spec, slot, res, rets, mon):
    """Draw a single fund cell in the 3×2 grid."""
    row, col = divmod(slot, 2)
    inner = gridspec.GridSpecFromSubplotSpec(
        1, 1, subplot_spec=outer_spec[row, col])
    ax = fig.add_subplot(inner[0])

    h_val  = res.get("h_threshold", 19.46)
    L_now  = res.get("L_now", 0.0)
    alarm  = res.get("alarm", False)
    in_alarm_state = L_now >= h_val
    in_watch_state = L_now >= 0.75 * h_val and not in_alarm_state

    # Cell background tint
    if in_alarm_state:
        cell_bg = "#FEF2F0"
    elif in_watch_state:
        cell_bg = "#FFFBEB"
    else:
        cell_bg = WHITE

    ax.set_facecolor(cell_bg)

    # ── Draw last-60-month CUSUM chart ────────────────────────────────────────
    df = pd.DataFrame(mon.history) if not hasattr(mon, 'history_df') else mon.history_df
    if df is None or (hasattr(df, '__len__') and len(df) == 0):
        ax.text(0.5, 0.5, "No data", ha="center", va="center",
                transform=ax.transAxes, color="grey")
        ax.axis("off")
        return

    # Use last 60 months
    CHART_MONTHS = 60
    n_total = len(df)
    start_idx = max(0, n_total - CHART_MONTHS)
    df_chart = df.iloc[start_idx:].copy()
    df_chart = df_chart.reset_index(drop=True)

    x = np.arange(len(df_chart))
    L_vals = df_chart["L"].values if "L" in df_chart.columns else np.zeros(len(df_chart))

    # Shade calibration window if visible in the 60m chart
    calib_end_global = CALIB_MONTHS
    calib_end_in_chart = calib_end_global - start_idx
    if 0 < calib_end_in_chart < len(df_chart):
        ax.axvspan(0, calib_end_in_chart, alpha=0.12, color="grey",
                   label="calibration")

    ax.plot(x, L_vals, lw=1.0, color=NAVY)
    ax.axhline(h_val, color=CORAL, ls="--", lw=0.9, alpha=0.9)
    ax.text(len(x) * 0.98, h_val * 1.02, f"h={h_val:.1f}",
            fontsize=5.5, color=CORAL, ha="right", va="bottom")

    # Alarm markers
    for at_global in mon.all_alarms:
        at_local = at_global - start_idx
        if 0 <= at_local < len(df_chart):
            ax.axvline(at_local, color=CORAL, lw=0.8, alpha=0.7)

    # Inverted y-axis: L=0 at top, L=h at bottom
    ax.invert_yaxis()
    ax.set_ylim(h_val * 1.15, 0)

    # x-axis: year labels from dates
    if rets is not None and len(rets) > 0:
        dates_chart = rets.index[start_idx:start_idx + len(df_chart)]
        if len(dates_chart) > 0:
            yr_ticks = []
            yr_labels = []
            seen_years = set()
            for i2, dt in enumerate(dates_chart):
                yr = dt.year
                if yr not in seen_years:
                    yr_ticks.append(i2)
                    yr_labels.append(str(yr))
                    seen_years.add(yr)
            ax.set_xticks(yr_ticks)
            ax.set_xticklabels(yr_labels, fontsize=5, rotation=45)
    else:
        ax.set_xticks([])

    ax.tick_params(axis="y", labelsize=5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # ── Stats text block (title + stats above chart) ──────────────────────────
    # We use the axes title area and annotations
    ir      = res.get("ir", float("nan"))
    ir_lo   = res.get("ir_lo", float("nan"))
    ir_hi   = res.get("ir_hi", float("nan"))
    te      = res.get("te", float("nan"))
    hit     = res.get("hit_rate")
    up_cap  = res.get("up_cap")
    dn_cap  = res.get("dn_cap")
    months  = res.get("months", 0)
    n_alm   = res.get("n_alarms", 0)
    alm_dt  = res.get("alarm_date", "-")
    vix_reg = res.get("vix_regime", "—")
    dq_flags= res.get("dq_flags", [])

    # Title: fund ticker + name truncated
    name_short = res.get("name", "")
    if len(name_short) > 35:
        name_short = name_short[:32] + "..."
    title_col = CORAL if in_alarm_state else (AMBER if in_watch_state else NAVY)
    ax.set_title(f"{res['fund']}  vs  {res['bench']}", fontsize=7.5,
                 fontweight="bold", color=title_col, pad=2)

    # Stats annotation in upper-left
    ci_str  = (f"IR {ir:+.2f} [{ir_lo:+.2f},{ir_hi:+.2f}]"
               if not (np.isnan(ir_lo) or np.isnan(ir_hi)) else f"IR {ir:+.2f}")
    te_str  = f"TE {te:.1%}" if not np.isnan(te) else "TE n/a"
    hit_str = f"Hit {hit:.0%}" if hit is not None else ""
    cap_str = (f"Up/Dn {up_cap:.0%}/{dn_cap:.0%}"
               if up_cap is not None and dn_cap is not None
               and not (np.isnan(up_cap) or np.isnan(dn_cap)) else "")
    alm_str = (f"{n_alm} alarm(s), last {alm_dt}" if n_alm > 0 else "No alarms")
    reg_str = f"regime: {vix_reg}" if alm_dt != "-" else ""

    stats_lines = [
        f"{months}m  {ci_str}",
        f"{te_str}  {hit_str}",
        cap_str,
        f"{alm_str}  {reg_str}",
    ]
    stats_text = "\n".join(l for l in stats_lines if l.strip())

    alarm_label = "IN ALARM" if in_alarm_state else ("WATCH" if in_watch_state else "")
    if alarm_label:
        ax.text(0.02, 0.98, alarm_label, transform=ax.transAxes,
                fontsize=6, fontweight="bold", color=title_col,
                va="top", ha="left",
                bbox=dict(boxstyle="round,pad=0.2", facecolor=cell_bg,
                          edgecolor=title_col, alpha=0.9))

    ax.text(0.98, 0.98, stats_text, transform=ax.transAxes,
            fontsize=5.2, ha="right", va="top", color="#333333",
            linespacing=1.35,
            bbox=dict(boxstyle="round,pad=0.25", facecolor=WHITE,
                      edgecolor=BORDER, alpha=0.85))

    # DQ badge
    if dq_flags:
        ax.text(0.02, 0.06, "\u26a0 DQ", transform=ax.transAxes,
                fontsize=5.5, color=AMBER, fontweight="bold", va="bottom")

    ax.set_xlabel("", fontsize=0)
    ax.set_ylabel("L (0=top)", fontsize=5, color="grey")


def page_fund_grid(pdf, all_rets, all_monitors, all_results, pnum, total_pages):
    # Sort by asset class, then fund ticker
    sorted_idx = sorted(range(len(all_results)),
                        key=lambda i: (_ac_sort_key(
                            _asset_class(all_results[i]["fund"], all_results[i]["bench"])),
                            all_results[i]["fund"]))

    n_funds      = len(sorted_idx)
    n_grid_pages = max(1, math.ceil(n_funds / FUNDS_PER_GRID_PAGE))

    for pp in range(n_grid_pages):
        chunk_idx = sorted_idx[pp * FUNDS_PER_GRID_PAGE:(pp + 1) * FUNDS_PER_GRID_PAGE]
        n_in_chunk = len(chunk_idx)
        n_rows     = math.ceil(n_in_chunk / 2)

        fig = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.980,
                 f"Fund Universe Grid  [{pp+1} / {n_grid_pages}]"
                 f"   (last 60 months shown, L inverted: 0=top h=bottom)",
                 ha="center", fontsize=9.5, fontweight="bold", color=NAVY)

        outer = gridspec.GridSpec(n_rows, 2, figure=fig,
                                  left=0.06, right=0.97,
                                  top=0.965, bottom=0.05,
                                  hspace=0.65, wspace=0.28)

        for slot, idx in enumerate(chunk_idx):
            res = all_results[idx]
            rets = all_rets[idx] if idx < len(all_rets) else None
            mon  = all_monitors[idx] if idx < len(all_monitors) else None
            if mon is None:
                continue
            try:
                _draw_grid_cell(fig, outer, slot, res, rets, mon)
            except Exception as e:
                row2, col2 = divmod(slot, 2)
                try:
                    inner = gridspec.GridSpecFromSubplotSpec(
                        1, 1, subplot_spec=outer[row2, col2])
                    ax_err = fig.add_subplot(inner[0])
                    ax_err.text(0.5, 0.5, f"Error: {e}", ha="center", va="center",
                                transform=ax_err.transAxes, color="red", fontsize=7)
                    ax_err.axis("off")
                except Exception:
                    pass

        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE N+1 — Summary Results Table
# ═══════════════════════════════════════════════════════════════════════════════

def page_summary_table(pdf, all_results, pnum, total_pages):
    rows  = _build_grouped_rows_summary(all_results)
    pages = _paginate_grouped_rows(rows, SUMMARY_PER_PAGE)
    n_pages = len(pages)

    headers  = ["Fund", "Bench", "Asset Class", "Mo.", "IR", "t-stat",
                "TE", "Hit%", "Up", "Dn", "L now", "h", "#Alm",
                "Last Alarm", "Regime", "Signal", "Flags"]
    col_x    = [0.030, 0.083, 0.128, 0.255, 0.295, 0.345,
                0.390, 0.433, 0.473, 0.510, 0.548, 0.593,
                0.633, 0.668, 0.735, 0.790, 0.855]

    DATA_ROW_H = 0.030
    HDR_ROW_H  = 0.020

    for sp, chunk in enumerate(pages):
        fig = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.962, "Summary Results Table", ha="center",
                 fontsize=18, fontweight="bold", color=NAVY)
        fig.text(0.5, 0.932,
                 f"All 67 funds — one row each   [{sp+1}/{n_pages}]",
                 ha="center", fontsize=9, color="grey")

        fig.add_artist(FancyBboxPatch((0.025, 0.875), 0.95, 0.028,
                                      boxstyle="round,pad=0.003", facecolor=NAVY,
                                      edgecolor="none", transform=fig.transFigure))
        for hdr, x in zip(headers, col_x):
            fig.text(x, 0.884, hdr, fontsize=6, fontweight="bold",
                     color=WHITE, va="center")

        y = 0.860
        data_idx = 0
        for kind, data in chunk:
            if kind == "header":
                ac     = data
                ac_col = ASSET_CLASS_COLORS.get(ac, NAVY)
                fig.add_artist(FancyBboxPatch((0.025, y - HDR_ROW_H + 0.002),
                                              0.95, HDR_ROW_H - 0.003,
                                              boxstyle="round,pad=0.002",
                                              facecolor=ac_col, edgecolor="none",
                                              transform=fig.transFigure, alpha=0.12))
                fig.text(0.032, y - HDR_ROW_H * 0.45,
                         f"  \u25aa  {ac.upper()}",
                         fontsize=7, fontweight="bold", color=ac_col, va="center")
                y -= HDR_ROW_H
            else:
                r = data
                h_val = r.get("h_threshold", 19.46)
                L_now = r.get("L_now", 0.0)
                in_alarm_state = L_now >= h_val
                in_watch_state = L_now >= 0.75 * h_val and not in_alarm_state

                if in_alarm_state:
                    bg = "#FEE2E2"
                elif in_watch_state:
                    bg = "#FEF9C3"
                elif data_idx % 2 == 0:
                    bg = CARD
                else:
                    bg = WHITE

                y -= DATA_ROW_H
                fig.add_artist(FancyBboxPatch((0.025, y - 0.003), 0.95, DATA_ROW_H - 0.002,
                                              boxstyle="round,pad=0.002", facecolor=bg,
                                              edgecolor="none", transform=fig.transFigure))

                ir    = r.get("ir", float("nan"))
                ir_lo = r.get("ir_lo", float("nan"))
                ir_hi = r.get("ir_hi", float("nan"))
                t     = r.get("t_stat") or 0.0
                te    = r.get("te", float("nan"))
                hit   = r.get("hit_rate")
                uc    = r.get("up_cap")
                dc    = r.get("dn_cap")
                sq    = r.get("sig_quality")
                n_alm = r.get("n_alarms", 0)
                alm_dt= r.get("alarm_date", "-")
                vix_r = r.get("vix_regime", "—")
                dq    = r.get("dq_flags", [])
                ac    = _asset_class(r["fund"], r["bench"])

                # Signal quality: "pending" for alarms < 12 months old
                alarm_ts = r.get("alarm_date_ts")
                sq_str = "—"
                if alarm_ts is not None:
                    months_since = (pd.Timestamp.today() - alarm_ts).days / 30.44
                    if months_since < 12:
                        sq_str = "pending"
                    elif sq is not None:
                        sq_str = f"{sq:.0%}"
                elif sq is not None:
                    sq_str = f"{sq:.0%}"

                ir_col = TEAL if t > 1.65 else (CORAL if ir < 0 else NAVY)
                l_col  = CORAL if in_alarm_state else (AMBER if in_watch_state else NAVY)
                dq_str = f"{len(dq)} flag(s)" if dq else "—"

                vals = [
                    r["fund"], r["bench"], ac[:14],
                    str(r.get("months", "?")),
                    f"{ir:+.2f}" if not np.isnan(ir) else "n/a",
                    f"{t:+.2f}" if t else "—",
                    f"{te:.1%}" if not np.isnan(te) else "n/a",
                    f"{hit:.0%}" if hit is not None else "—",
                    f"{uc:.0%}" if uc is not None and not np.isnan(uc) else "—",
                    f"{dc:.0%}" if dc is not None and not np.isnan(dc) else "—",
                    f"{L_now:.1f}",
                    f"{h_val:.1f}",
                    str(n_alm) if n_alm else "—",
                    alm_dt if alm_dt != "-" else "—",
                    vix_r,
                    sq_str,
                    dq_str,
                ]
                colors = [
                    NAVY, NAVY, ASSET_CLASS_COLORS.get(ac, NAVY),
                    NAVY, ir_col, ir_col,
                    NAVY, NAVY, NAVY, NAVY,
                    l_col, NAVY,
                    CORAL if n_alm > 0 else NAVY,
                    NAVY, AMBER if vix_r in ("elevated","crisis") else NAVY,
                    AMBER if sq_str == "pending" else NAVY,
                    AMBER if dq else NAVY,
                ]

                for val, x, col in zip(vals, col_x, colors):
                    fig.text(x, y + DATA_ROW_H * 0.22, val, fontsize=5.8,
                             va="center", color=col)
                data_idx += 1

        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE N+2 — Alarm Log
# ═══════════════════════════════════════════════════════════════════════════════

def page_alarm_log(pdf, all_results, all_post_alarm, pnum, total_pages):
    # Build flat alarm events list (all alarms, all funds)
    alarm_events = []
    for res, pa_df in zip(all_results, all_post_alarm):
        if not res["alarm"]:
            continue
        mon_alarms = res.get("all_alarm_dates", [])
        h_val = res.get("h_threshold", 19.46)
        # For each alarm, find post-alarm data
        if pa_df is not None and not pa_df.empty:
            for _, grp in pa_df.groupby("alarm_t"):
                alarm_t   = grp["alarm_t"].iloc[0]
                alarm_dt  = grp["alarm_date"].iloc[0]
                row12 = grp[grp["horizon"] == 12]
                row24 = grp[grp["horizon"] == 24]
                exc12 = row12["excess_cum"].iloc[0] if not row12.empty else float("nan")
                exc24 = row24["excess_cum"].iloc[0] if not row24.empty else float("nan")
                correct12 = row12["correct_alarm"].iloc[0] if not row12.empty else None

                # Verdict
                if correct12 is None or (isinstance(exc12, float) and np.isnan(exc12)):
                    verdict = "pending"
                elif correct12:
                    verdict = "correct"
                else:
                    verdict = "false alarm"

                # L at crossing — approximate (threshold)
                alarm_events.append({
                    "date":     alarm_dt,
                    "fund":     res["fund"],
                    "bench":    res["bench"],
                    "L_cross":  h_val,
                    "h":        h_val,
                    "regime":   res.get("vix_regime", "—"),
                    "exc12":    exc12,
                    "exc24":    exc24,
                    "verdict":  verdict,
                    "action":   "—",
                })
        else:
            # Alarm exists but no pa_df
            alarm_events.append({
                "date":     res.get("alarm_date", "—"),
                "fund":     res["fund"],
                "bench":    res["bench"],
                "L_cross":  h_val,
                "h":        h_val,
                "regime":   res.get("vix_regime", "—"),
                "exc12":    float("nan"),
                "exc24":    float("nan"),
                "verdict":  "pending",
                "action":   "—",
            })

    # Sort most recent first
    def _sort_key(ev):
        try:
            return pd.Timestamp(ev["date"])
        except Exception:
            return pd.Timestamp("1900-01-01")

    alarm_events.sort(key=_sort_key, reverse=True)

    if not alarm_events:
        alarm_events = [{"date": "—", "fund": "(none)", "bench": "—",
                         "L_cross": 0, "h": 0, "regime": "—",
                         "exc12": float("nan"), "exc24": float("nan"),
                         "verdict": "—", "action": "—"}]

    headers = ["Date", "Fund", "Bench", "L at cross", "h in force",
               "VIX regime", "+12M outcome", "+24M outcome", "Verdict", "Action"]
    col_x   = [0.030, 0.098, 0.160, 0.215, 0.268, 0.323, 0.410, 0.510, 0.600, 0.700]

    n_pages = max(1, math.ceil(len(alarm_events) / ALARM_LOG_PER_PAGE))
    DATA_ROW_H = 0.042

    for ap in range(n_pages):
        chunk = alarm_events[ap * ALARM_LOG_PER_PAGE:(ap + 1) * ALARM_LOG_PER_PAGE]
        fig   = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.962, "Alarm Log — Audit Trail", ha="center",
                 fontsize=18, fontweight="bold", color=NAVY)
        fig.text(0.5, 0.932,
                 f"All alarm events, most recent first   [{ap+1}/{n_pages}]",
                 ha="center", fontsize=9, color="grey")

        fig.add_artist(FancyBboxPatch((0.025, 0.878), 0.95, 0.028,
                                      boxstyle="round,pad=0.003", facecolor=NAVY,
                                      edgecolor="none", transform=fig.transFigure))
        for hdr, x in zip(headers, col_x):
            fig.text(x, 0.886, hdr, fontsize=7, fontweight="bold",
                     color=WHITE, va="center")

        y = 0.858
        for r_idx, ev in enumerate(chunk):
            bg   = CARD if r_idx % 2 == 0 else WHITE
            y   -= DATA_ROW_H
            fig.add_artist(FancyBboxPatch((0.025, y - 0.005), 0.95, DATA_ROW_H - 0.003,
                                          boxstyle="round,pad=0.002", facecolor=bg,
                                          edgecolor="none", transform=fig.transFigure))

            exc12 = ev["exc12"]
            exc24 = ev["exc24"]
            verdict = ev["verdict"]

            exc12_str = f"{exc12:+.1%}" if not np.isnan(exc12) else "pending"
            exc24_str = f"{exc24:+.1%}" if not np.isnan(exc24) else "pending"
            v_col     = (TEAL if verdict == "correct" else
                         (CORAL if verdict == "false alarm" else AMBER))
            reg_col   = (CORAL if ev["regime"] == "crisis" else
                         (AMBER if ev["regime"] == "elevated" else TEAL))

            vals   = [ev["date"], ev["fund"], ev["bench"],
                      f"{ev['L_cross']:.2f}", f"{ev['h']:.2f}",
                      ev["regime"], exc12_str, exc24_str, verdict, ev["action"]]
            colors = [NAVY, NAVY, NAVY,
                      CORAL, NAVY,
                      reg_col,
                      CORAL if not np.isnan(exc12) and exc12 < 0 else TEAL,
                      CORAL if not np.isnan(exc24) and exc24 < 0 else TEAL,
                      v_col, NAVY]

            for val, x, col in zip(vals, col_x, colors):
                fig.text(x, y + DATA_ROW_H * 0.22, val, fontsize=7.5,
                         va="center", color=col)

        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE N+3 — Framework Health Diagnostics
# ═══════════════════════════════════════════════════════════════════════════════

def page_framework_health(pdf, all_results, all_post_alarm, pnum, total_pages):
    fig = plt.figure(figsize=(11, 8.5))
    fig.patch.set_facecolor(WHITE)
    fig.text(0.5, 0.962, "Framework Health Diagnostics", ha="center",
             fontsize=18, fontweight="bold", color=NAVY)
    fig.text(0.5, 0.932, "Signal quality · alarm distribution · h sensitivity",
             ha="center", fontsize=10, color="grey")

    # ── Panel 1: Signal quality by regime bar chart ───────────────────────────
    ax1 = fig.add_axes([0.06, 0.62, 0.38, 0.27])
    ax1.set_facecolor(CARD)

    regimes_display = ["calm", "elevated", "crisis"]
    regime_internal = {"calm": "low_vol", "elevated": "med_vol", "crisis": "high_vol"}
    horizons = [12, 24]

    sq_data = {h: {} for h in horizons}
    for h in horizons:
        for rd in regimes_display:
            ri = regime_internal[rd]
            sq_vals = []
            for res, pa_df in zip(all_results, all_post_alarm):
                if pa_df is None or pa_df.empty:
                    continue
                # filter to this regime
                rows = pa_df[pa_df["horizon"] == h].dropna(subset=["correct_alarm"])
                if "regime" in rows.columns:
                    rows = rows[rows["regime"] == ri]
                if len(rows) > 0:
                    sq_vals.append(float(rows["correct_alarm"].mean()))
            sq_data[h][rd] = float(np.mean(sq_vals)) if sq_vals else 0.0

    x_pos = np.arange(len(regimes_display))
    width = 0.3
    bars12 = ax1.bar(x_pos - width/2,
                     [sq_data[12].get(r, 0) for r in regimes_display],
                     width, label="+12M", color=NAVY, alpha=0.8)
    bars24 = ax1.bar(x_pos + width/2,
                     [sq_data[24].get(r, 0) for r in regimes_display],
                     width, label="+24M", color=TEAL, alpha=0.8)

    ax1.axhline(0.60, color=CORAL, ls="--", lw=1.2, label="Target 60%")
    ax1.set_xticks(x_pos)
    ax1.set_xticklabels(regimes_display, fontsize=8)
    ax1.set_ylabel("Signal quality", fontsize=7)
    ax1.set_ylim(0, 1.0)
    ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0%}"))
    ax1.set_title("Signal Quality by Regime", fontsize=9, fontweight="bold", color=NAVY)
    ax1.legend(fontsize=7)
    ax1.tick_params(labelsize=7)
    for bar in list(bars12) + list(bars24):
        h2 = bar.get_height()
        if h2 > 0:
            ax1.text(bar.get_x() + bar.get_width()/2, h2 + 0.01,
                     f"{h2:.0%}", ha="center", va="bottom", fontsize=6)

    # ── Panel 2: Alarm distribution by year ───────────────────────────────────
    ax2 = fig.add_axes([0.54, 0.62, 0.40, 0.27])
    ax2.set_facecolor(CARD)

    alarm_years = {}
    for res in all_results:
        ts = res.get("alarm_date_ts")
        if ts is not None:
            yr = ts.year
            alarm_years[yr] = alarm_years.get(yr, 0) + 1

    if alarm_years:
        years  = sorted(alarm_years.keys())
        counts = [alarm_years[y] for y in years]
        bar_cols = [CORAL if c > 3 else AMBER for c in counts]
        ax2.bar(years, counts, color=bar_cols, edgecolor="white", lw=0.5)
        ax2.set_xticks(years)
        ax2.set_xticklabels([str(y) for y in years], fontsize=6.5, rotation=45)
        yr_labels = "  ".join(f"{y}: {c}" for y, c in zip(years, counts))
        ax2.text(0.5, -0.22, yr_labels, transform=ax2.transAxes,
                 ha="center", fontsize=6, color="#444444")
    else:
        ax2.text(0.5, 0.5, "No alarms recorded", ha="center", va="center",
                 transform=ax2.transAxes, color="grey")

    ax2.set_title("Alarm Distribution by Year", fontsize=9, fontweight="bold", color=NAVY)
    ax2.set_ylabel("Number of alarms", fontsize=7)
    ax2.tick_params(labelsize=7)

    # ── Panel 3: h sensitivity ────────────────────────────────────────────────
    ax3 = fig.add_axes([0.06, 0.28, 0.88, 0.25])
    ax3.set_facecolor(CARD)
    ax3.axis("off")

    ax3.text(0.5, 0.95, "h Sensitivity Analysis", ha="center", va="top",
             fontsize=10, fontweight="bold", color=NAVY, transform=ax3.transAxes)

    # Count alarms at h, h-2, h+2 per asset class
    sensitivity_rows = []
    for preset_key, preset_vals in ASSET_CLASS_PRESETS.items():
        h_base = preset_vals.get("threshold", 19.46)
        h_lo   = h_base - 2.0
        h_hi   = h_base + 2.0
        # Funds in this preset
        n_base = sum(1 for r in all_results
                     if _preset_for(r["fund"], r["bench"]) == preset_key and r["alarm"])
        # Approximate: at h-2 alarms increase, at h+2 decrease
        # We use L_now to estimate: alarms at h-2 = funds where max_L >= h-2
        n_lo_extra  = sum(1 for r in all_results
                          if _preset_for(r["fund"], r["bench"]) == preset_key
                          and not r["alarm"]
                          and r.get("L_now", 0) >= h_lo)
        n_hi_fewer  = sum(1 for r in all_results
                          if _preset_for(r["fund"], r["bench"]) == preset_key
                          and r["alarm"]
                          and r.get("L_now", 0) < h_hi)
        sensitivity_rows.append((preset_key, h_base, h_lo, h_hi, n_base, n_lo_extra, n_hi_fewer))

    col_sens = [0.02, 0.15, 0.27, 0.40, 0.53, 0.67, 0.82]
    hdrs_s   = ["Asset class", "h (base)", "h-2", "h+2",
                "Alarms @ h", "Extra @ h-2", "Fewer @ h+2"]
    for hdr, x in zip(hdrs_s, col_sens):
        ax3.text(x, 0.78, hdr, fontsize=7.5, fontweight="bold", color=NAVY,
                 transform=ax3.transAxes)

    for i, (pk, hb, hl, hh, nb, nl, nh) in enumerate(sensitivity_rows):
        yrow = 0.60 - i * 0.15
        ax3.text(col_sens[0], yrow, pk,        fontsize=7, transform=ax3.transAxes)
        ax3.text(col_sens[1], yrow, f"{hb:.2f}", fontsize=7, transform=ax3.transAxes)
        ax3.text(col_sens[2], yrow, f"{hl:.2f}", fontsize=7, transform=ax3.transAxes)
        ax3.text(col_sens[3], yrow, f"{hh:.2f}", fontsize=7, transform=ax3.transAxes)
        ax3.text(col_sens[4], yrow, str(nb),   fontsize=7, transform=ax3.transAxes)
        ax3.text(col_sens[5], yrow, f"+{nl}",  fontsize=7, color=CORAL,
                 transform=ax3.transAxes)
        ax3.text(col_sens[6], yrow, f"-{nh}",  fontsize=7, color=TEAL,
                 transform=ax3.transAxes)

    # ── Health summary text ───────────────────────────────────────────────────
    ax4 = fig.add_axes([0.06, 0.05, 0.88, 0.18])
    ax4.set_facecolor("#EBF5FB")
    for s in ax4.spines.values():
        s.set_visible(False)
    ax4.set_xticks([]); ax4.set_yticks([])

    calm_sq_vals = [sq_data[12].get("calm", None)]
    calm_q = calm_sq_vals[0]
    total_alm = sum(1 for r in all_results if r["alarm"])
    n_funds   = len(all_results)

    sq_label = "(TARGET MET \u2713)" if calm_q and calm_q >= 0.60 else "(BELOW TARGET \u2014 REVIEW CALIBRATION)"
    sq_pct   = f"{calm_q:.0%}" if calm_q is not None else "n/a"
    health_text = (
        f"  Total alarms: {total_alm} / {n_funds} funds ({total_alm/max(n_funds,1):.0%})    "
        f"Calm-regime signal quality +12M: {sq_pct}  {sq_label}    "
        f"Expected false alarms/year at current h: ~{n_funds * 12 / 60:.1f}"
    )
    ax4.text(0.5, 0.65, health_text, ha="center", va="center",
             fontsize=8, color="#333333", transform=ax4.transAxes, linespacing=1.5)
    ax4.text(0.5, 0.25,
             "Signal quality = % of alarms where fund underperformed benchmark at +12M horizon. "
             "Target \u2265 60% in calm regimes. Crisis alarms are expected to be less reliable "
             "(market-driven, not skill loss).",
             ha="center", va="center", fontsize=7.5, color="#555555",
             transform=ax4.transAxes, linespacing=1.4)

    page_footer(fig, pnum[0], total_pages)
    pnum[0] += 1
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE — Waterfall Chart (kept from original)
# ═══════════════════════════════════════════════════════════════════════════════

def page_waterfall(pdf, all_results, pnum, total_pages):
    fig = plt.figure(figsize=(11, 8.5))
    fig.patch.set_facecolor(WHITE)
    fig.text(0.5, 0.962, "Cumulative Excess Returns — Ranked", ha="center",
             fontsize=20, fontweight="bold", color=NAVY)
    fig.text(0.5, 0.930, "Total return of each fund vs its benchmark (full history)",
             ha="center", fontsize=11, color="grey")

    ax = fig.add_axes([0.24, 0.05, 0.69, 0.86])
    sorted_res = sorted(all_results, key=lambda x: x["cum_excess"])
    labels = [f"{r['fund']} vs {r['bench']}" for r in sorted_res]
    values = [r["cum_excess"] for r in sorted_res]
    colors = [TEAL if v >= 0 else CORAL for v in values]
    alarm  = [r["alarm"] for r in sorted_res]

    bar_h = max(0.3, min(0.7, 20 / max(len(sorted_res), 1)))
    bars  = ax.barh(labels, values, color=colors, height=bar_h,
                    edgecolor="white", lw=0.4)

    fs = max(5.5, min(8, 200 / max(len(sorted_res), 1)))
    for bar, val, alm in zip(bars, values, alarm):
        x_pos = val + 0.003 if val >= 0 else val - 0.003
        ha    = "left" if val >= 0 else "right"
        label = f"{val:+.1%}" + (" *" if alm else "")
        ax.text(x_pos, bar.get_y() + bar.get_height() / 2,
                label, va="center", ha=ha, fontsize=fs,
                color=TEAL if val >= 0 else CORAL, fontweight="bold")

    ax.axvline(0, color=NAVY, lw=1.2)
    ax.set_xlabel("Cumulative excess return vs benchmark", fontsize=10, color=NAVY)
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:.0%}"))
    ax.tick_params(axis="y", labelsize=max(5, int(fs)))
    ax.tick_params(axis="x", labelsize=9)
    ax.set_facecolor(LGREY)
    ax.grid(axis="x", alpha=0.4, lw=0.6)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    legend_elements = [
        mpatches.Patch(facecolor=TEAL,  label="Outperformed benchmark"),
        mpatches.Patch(facecolor=CORAL, label="Underperformed benchmark"),
        plt.Line2D([0], [0], color="none", label="* = CUSUM alarm triggered"),
    ]
    ax.legend(handles=legend_elements, loc="lower right", fontsize=9, framealpha=0.9)

    page_footer(fig, pnum[0], total_pages)
    pnum[0] += 1
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE — Insufficient History
# ═══════════════════════════════════════════════════════════════════════════════

def page_insufficient_history(pdf, insufficient_history, pnum, total_pages):
    PER_PAGE = 20
    if not insufficient_history:
        insufficient_history = []

    n_pages = max(1, math.ceil(max(len(insufficient_history), 1) / PER_PAGE))

    for pp in range(n_pages):
        chunk = insufficient_history[pp * PER_PAGE:(pp + 1) * PER_PAGE]
        fig   = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.962, "Insufficient History — Excluded from Main Analysis",
                 ha="center", fontsize=16, fontweight="bold", color=NAVY)
        fig.text(0.5, 0.932,
                 f"Funds with < {MIN_MONITORING_MONTHS} months of live CUSUM history   [{pp+1}/{n_pages}]",
                 ha="center", fontsize=9, color="grey")

        headers  = ["Ticker", "Fund Name", "Benchmark", "Months Available", "Note"]
        col_x    = [0.040, 0.130, 0.540, 0.640, 0.730]
        DATA_ROW_H = 0.040

        fig.add_artist(FancyBboxPatch((0.030, 0.875), 0.94, 0.028,
                                      boxstyle="round,pad=0.003", facecolor=NAVY,
                                      edgecolor="none", transform=fig.transFigure))
        for hdr, x in zip(headers, col_x):
            fig.text(x, 0.883, hdr, fontsize=8, fontweight="bold",
                     color=WHITE, va="center")

        y = 0.855
        if not chunk:
            fig.text(0.5, 0.5, "No funds with insufficient history.", ha="center",
                     va="center", fontsize=11, color=TEAL)
        else:
            for i, rec in enumerate(chunk):
                bg = CARD if i % 2 == 0 else WHITE
                y -= DATA_ROW_H
                fig.add_artist(FancyBboxPatch((0.030, y - 0.005), 0.94, DATA_ROW_H - 0.004,
                                              boxstyle="round,pad=0.002", facecolor=bg,
                                              edgecolor="none", transform=fig.transFigure))
                n_m   = rec.get("months", "?")
                note  = (f"Only {n_m}m available; CUSUM asymptotic theory requires "
                         f"\u2265{MIN_MONITORING_MONTHS}m live monitoring.")
                vals  = [rec.get("ticker", rec.get("fund", "?")),
                         rec.get("name", "—"),
                         rec.get("bench", "—"),
                         str(n_m),
                         note]
                cols  = [NAVY, NAVY, NAVY, AMBER, "#555555"]
                for val, x, col in zip(vals, col_x, cols):
                    fig.text(x, y + DATA_ROW_H * 0.22, val, fontsize=7.5,
                             va="center", color=col)

        insuff_note = (
            "These funds are monitored but excluded from aggregate statistics.\n"
            f"As history accumulates to \u2265{MIN_MONITORING_MONTHS} months they will be promoted to the main analysis."
        )
        fig.text(0.5, 0.090, insuff_note,
                 ha="center", va="top", fontsize=8.5, color="#444444", linespacing=1.5)

        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE — Data Quality Appendix
# ═══════════════════════════════════════════════════════════════════════════════

def page_data_quality(pdf, all_results, pnum, total_pages):
    dq_funds = [r for r in all_results if r.get("dq_flags")]
    PER_PAGE = 16

    if not dq_funds:
        fig = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.962, "Data Quality Appendix", ha="center",
                 fontsize=18, fontweight="bold", color=NAVY)
        fig.text(0.5, 0.500, "No data quality flags raised across the universe.",
                 ha="center", va="center", fontsize=14, color=TEAL)
        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)
        return

    n_pages = max(1, math.ceil(len(dq_funds) / PER_PAGE))
    for pp in range(n_pages):
        chunk = dq_funds[pp * PER_PAGE:(pp + 1) * PER_PAGE]
        fig   = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.962, "Data Quality Appendix", ha="center",
                 fontsize=18, fontweight="bold", color=NAVY)
        fig.text(0.5, 0.932,
                 f"Funds with one or more data quality flags   [{pp+1}/{n_pages}]",
                 ha="center", fontsize=9, color="grey")

        headers  = ["Fund", "Bench", "Asset Class", "# Flags", "Warnings"]
        col_x    = [0.030, 0.085, 0.130, 0.230, 0.275]
        DATA_ROW_H = 0.050

        fig.add_artist(FancyBboxPatch((0.025, 0.875), 0.95, 0.028,
                                      boxstyle="round,pad=0.003", facecolor=NAVY,
                                      edgecolor="none", transform=fig.transFigure))
        for hdr, x in zip(headers, col_x):
            fig.text(x, 0.883, hdr, fontsize=8, fontweight="bold",
                     color=WHITE, va="center")

        y = 0.855
        for i, r in enumerate(chunk):
            bg    = CARD if i % 2 == 0 else WHITE
            y    -= DATA_ROW_H
            flags = r.get("dq_flags", [])
            n_fl  = len(flags)
            ac    = _asset_class(r["fund"], r["bench"])

            fig.add_artist(FancyBboxPatch((0.025, y - 0.005), 0.95, DATA_ROW_H - 0.004,
                                          boxstyle="round,pad=0.002", facecolor=bg,
                                          edgecolor="none", transform=fig.transFigure))

            warn_txt = flags[0][:80] if flags else "—"
            if len(flags) > 1:
                warn_txt += f"  (+{len(flags)-1} more)"

            vals = [r["fund"], r["bench"], ac[:16], str(n_fl), warn_txt]
            cols = [NAVY, NAVY, ASSET_CLASS_COLORS.get(ac, NAVY), AMBER, "#555555"]
            for val, x, col in zip(vals, col_x, cols):
                fig.text(x, y + DATA_ROW_H * 0.28, val, fontsize=7.2,
                         va="center", color=col)

            # Show additional flags below
            if len(flags) > 1:
                for fi, flag in enumerate(flags[1:3], 1):
                    fig.text(col_x[4], y + DATA_ROW_H * 0.28 - fi * 0.012,
                             f"  {flag[:80]}", fontsize=6.5, color="#666666")

        fig.text(0.5, 0.045,
                 "DQ flags: stale NAV (3+ consecutive identical returns) | TE > 25% | "
                 "excess-return autocorrelation > 0.3 | |monthly return| > 50% | missing months",
                 ha="center", fontsize=7.5, color="#555555", style="italic")

        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE — Conclusions & Action Items
# ═══════════════════════════════════════════════════════════════════════════════

def page_conclusion(pdf, all_results, all_post_alarm, pnum, total_pages):
    today    = datetime.today()
    n_funds  = len(all_results)
    n_alarm  = sum(1 for r in all_results if r["alarm"])
    n_clean  = n_funds - n_alarm

    in_alarm_now = [r for r in all_results
                    if r.get("L_now", 0) >= r.get("h_threshold", 19.46)]
    watchlist    = [r for r in all_results
                    if r.get("L_now", 0) >= 0.75 * r.get("h_threshold", 19.46)
                    and r not in in_alarm_now]
    clean_funds  = [r for r in all_results if not r["alarm"]]

    headers = ["Fund", "Bench", "Asset Class", "IR", "t-stat",
               "#Alarms", "Last Alarm", "L now", "h", "State"]
    col_x   = [0.030, 0.090, 0.148, 0.300, 0.356, 0.408,
               0.460, 0.560, 0.612, 0.658]
    ROW_H   = 0.038

    alarmed_sorted = sorted(in_alarm_now + watchlist,
                            key=lambda r: r.get("L_now", 0), reverse=True)
    if not alarmed_sorted:
        alarmed_sorted = [r for r in all_results if r["alarm"]]

    n_concl_pages = max(1, math.ceil(max(len(alarmed_sorted), 1) / CONCLUSION_PER_PAGE))

    for cp in range(n_concl_pages):
        chunk   = alarmed_sorted[cp * CONCLUSION_PER_PAGE:(cp + 1) * CONCLUSION_PER_PAGE]
        is_last = (cp == n_concl_pages - 1)
        pg_lbl  = f"  —  {cp+1}/{n_concl_pages}" if n_concl_pages > 1 else ""

        fig = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor(WHITE)
        fig.text(0.5, 0.962, "Conclusions & Action Items", ha="center",
                 fontsize=20, fontweight="bold", color=NAVY)
        fig.text(0.5, 0.930,
                 f"Funds requiring attention as of {today.strftime('%d %B %Y')}{pg_lbl}",
                 ha="center", fontsize=10, color="grey")

        fig.add_artist(FancyBboxPatch((0.025, 0.875), 0.95, 0.028,
                                      boxstyle="round,pad=0.003", facecolor=NAVY,
                                      edgecolor="none", transform=fig.transFigure))
        for hdr, x in zip(headers, col_x):
            fig.text(x, 0.883, hdr, fontsize=7.5, fontweight="bold",
                     color=WHITE, va="center")

        y = 0.855
        for r_idx, r in enumerate(chunk):
            in_al = r.get("L_now", 0) >= r.get("h_threshold", 19.46)
            in_wt = r.get("L_now", 0) >= 0.75 * r.get("h_threshold", 19.46)
            state_str = "IN ALARM" if in_al else ("WATCHLIST" if in_wt else "ALARM")
            s_col     = CORAL if in_al else AMBER
            bg        = "#FEF2F0" if in_al else "#FFFBEB"

            ir   = r.get("ir", float("nan"))
            t    = r.get("t_stat") or 0.0
            L    = r.get("L_now", 0.0)
            h    = r.get("h_threshold", 19.46)
            ac   = _asset_class(r["fund"], r["bench"])

            y -= ROW_H
            fig.add_artist(FancyBboxPatch((0.025, y - 0.005), 0.95, ROW_H - 0.004,
                                          boxstyle="round,pad=0.002", facecolor=bg,
                                          edgecolor="none", transform=fig.transFigure))

            vals   = [r["fund"], r["bench"], ac[:14],
                      f"{ir:+.2f}" if not np.isnan(ir) else "n/a",
                      f"{t:+.1f}",
                      str(r.get("n_alarms", "—")),
                      r.get("alarm_date", "—"),
                      f"{L:.2f}", f"{h:.2f}", state_str]
            colors = [NAVY, NAVY, ASSET_CLASS_COLORS.get(ac, NAVY),
                      TEAL if not np.isnan(ir) and ir > 0 else CORAL,
                      NAVY, NAVY, NAVY,
                      s_col, NAVY, s_col]

            for val, x, col in zip(vals, col_x, colors):
                fw = "bold" if col in (CORAL, AMBER) else "normal"
                fig.text(x, y + ROW_H * 0.22, val, fontsize=7.5,
                         va="center", color=col, fontweight=fw)

        if is_last:
            bottom_y = y - ROW_H * 0.5

            if clean_funds:
                clean_line = ", ".join(r["fund"] for r in clean_funds[:10])
                if len(clean_funds) > 10:
                    clean_line += f" ... +{len(clean_funds)-10} more"
                cln_y = max(0.300, bottom_y - 0.008)
                fig.text(0.030, cln_y,
                         f"No alarm ({len(clean_funds)} funds):  {clean_line}",
                         fontsize=7.5, color=TEAL, fontstyle="italic")
                bottom_y = cln_y - 0.038

            box_y = 0.048
            box_h = min(0.155, max(0.080, bottom_y - box_y - 0.005))

            ax_cav = fig.add_axes([0.030, box_y, 0.455, box_h])
            ax_cav.set_facecolor("#FEF2F0")
            for s in ax_cav.spines.values():
                s.set_visible(False)
            ax_cav.set_xticks([]); ax_cav.set_yticks([])
            ax_cav.text(0.5, 0.96, "Statistical Caveats",
                        ha="center", va="top", fontsize=9, fontweight="bold",
                        color=CORAL, transform=ax_cav.transAxes)
            exp_fa_str = f"{n_funds*12/60:.1f}"
            caveats = (
                f"  \u2022  Multiple testing: {n_funds} funds \u00d7 ARL=60m "
                "\u2192 ~" + exp_fa_str + " expected false alarms/year.\n"
                "  \u2022  Regime alarms (crisis/elevated VIX): likely market-driven,\n"
                "     not skill loss. See Alarm Log for regime breakdown.\n"
                "  \u2022  An alarm triggers a structured REVIEW, NOT auto-termination."
            )
            ax_cav.text(0.02, 0.74, caveats, va="top", fontsize=7.8,
                        color="#333333", transform=ax_cav.transAxes, linespacing=1.55)

            ax_next = fig.add_axes([0.515, box_y, 0.455, box_h])
            ax_next.set_facecolor("#EBF5FB")
            for s in ax_next.spines.values():
                s.set_visible(False)
            ax_next.set_xticks([]); ax_next.set_yticks([])
            ax_next.text(0.5, 0.96, "Recommended Next Steps",
                         ha="center", va="top", fontsize=9, fontweight="bold",
                         color=NAVY, transform=ax_next.transAxes)
            n_alarm_str   = str(len(in_alarm_now))
            n_watch_str   = str(len(watchlist))
            steps = (
                "  \u2022  Schedule IC review for all " + n_alarm_str + " IN ALARM fund(s).\n"
                "  \u2022  Place " + n_watch_str + " WATCHLIST fund(s) on heightened monitoring.\n"
                "  \u2022  Verify DQ-flagged funds against official NAV sources.\n"
                "  \u2022  Re-run report monthly after each month-end NAV update."
            )
            ax_next.text(0.02, 0.74, steps, va="top", fontsize=7.8,
                         color="#333333", transform=ax_next.transAxes, linespacing=1.55)

        page_footer(fig, pnum[0], total_pages)
        pnum[0] += 1
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════════════
# DATA COLLECTION
# ═══════════════════════════════════════════════════════════════════════════════

class MonitorResult:
    """Thin wrapper to attach history_df to a CUSUMMonitor for downstream use."""
    def __init__(self, monitor, history_df):
        for k, v in vars(monitor).items():
            setattr(self, k, v)
        self.history_df = history_df
        # also expose .history list via history_df for _draw_grid_cell
        if history_df is not None:
            self.history = history_df.to_dict("records")


def collect_data():
    """
    Fetch and run CUSUM for every fund in PRESET_PAIRS.

    Returns
    -------
    all_rets, all_monitors, all_results, all_post_alarm,
    insufficient_history, vix_series
    """
    # ── Verify no duplicate tickers ──────────────────────────────────────────
    tickers = list(PRESET_PAIRS.keys())
    assert len(set(tickers)) == len(tickers), \
        f"Duplicate tickers in PRESET_PAIRS: {[t for t in tickers if tickers.count(t)>1]}"

    # ── Load VIX ──────────────────────────────────────────────────────────────
    vix_series = None
    try:
        vix_series = load_vix_history()
        print(f"  VIX loaded: {len(vix_series)} monthly observations")
    except (ImportError, RuntimeError, Exception) as e:
        print(f"  VIX download failed ({e}); regime labels will be 'unknown'")

    all_rets           = []
    all_monitors       = []
    all_results        = []
    all_post_alarm     = []
    insufficient_history = []

    print("Fetching data from Yahoo Finance...")
    for ticker, (fund, bench, name) in PRESET_PAIRS.items():
        print(f"  {ticker}...", end=" ", flush=True)
        try:
            rets = load_from_yfinance(fund, bench)

            if rets is None or len(rets) == 0:
                print("EMPTY — skipped")
                continue

            # DQ check
            dq_flags = []
            try:
                dq_flags = sanity_check_returns(rets, fund_name=fund)
            except Exception as dq_e:
                dq_flags = [f"DQ check failed: {dq_e}"]

            calib    = rets.iloc[:CALIB_MONTHS]
            excess_s = calib["portfolio"] - calib["benchmark"]
            init_te  = max(excess_s.std() * (12 ** 0.5), 0.01)

            ac      = _asset_class(fund, bench)
            preset  = AC_TO_PRESET.get(ac, "equity")
            monitor = CUSUMMonitor(initial_te=init_te, preset=preset,
                                   calibration_months=CALIB_MONTHS,
                                   rolling_restart=True)
            history_df = monitor.run(rets)

            n_months   = len(rets)
            mon_months = n_months - CALIB_MONTHS
            h_val      = monitor.threshold

            excess_log = np.log((1 + rets["portfolio"]) / (1 + rets["benchmark"]))
            ir  = (excess_log.mean() * 12) / (excess_log.std() * (12 ** 0.5)) \
                  if excess_log.std() > 0 else 0.0
            te  = excess_log.std() * (12 ** 0.5)
            cum = (1 + rets["portfolio"]).prod() / (1 + rets["benchmark"]).prod() - 1

            try:
                ir_pt, ir_lo, ir_hi = bootstrap_ir_ci(rets)
            except Exception:
                ir_pt, ir_lo, ir_hi = ir, float("nan"), float("nan")

            extra = {}
            try:
                extra = compute_extra_stats(rets)
            except Exception:
                pass

            # Alarm date (first alarm)
            alarm_date = (rets.index[monitor.alarm_t - 1].strftime("%b %Y")
                          if monitor.alarm and monitor.alarm_t else "-")
            alarm_date_ts = (rets.index[monitor.alarm_t - 1]
                             if monitor.alarm and monitor.alarm_t else None)

            # All alarm dates (for alarm log)
            all_alarm_dates = []
            for at in monitor.all_alarms:
                if at <= len(rets):
                    all_alarm_dates.append(rets.index[at - 1].strftime("%Y-%m"))

            # VIX regime
            vix_regime = "—"
            if alarm_date != "-" and vix_series is not None:
                try:
                    reg_map   = tag_alarm_regime([alarm_date], vix_series)
                    raw_label = reg_map.get(alarm_date, "unknown")
                    vix_regime = VIX_REGIME_DISPLAY.get(raw_label, raw_label)
                except Exception:
                    vix_regime = "unknown"
            elif alarm_date != "-":
                vix_regime = "unknown"

            pa_df = post_alarm_returns(rets, monitor)
            sig_q = signal_quality(pa_df, horizon=12)

            short_hist = mon_months < MIN_MONITORING_MONTHS

            mr = MonitorResult(monitor, history_df)

            res = {
                "fund":            fund,
                "bench":           bench,
                "name":            name,
                "months":          n_months,
                "mon_months":      mon_months,
                "short_hist":      short_hist,
                "ir":              ir,
                "ir_lo":           ir_lo,
                "ir_hi":           ir_hi,
                "te":              te,
                "cum_excess":      cum,
                "hit_rate":        extra.get("hit_rate"),
                "t_stat":          extra.get("t_stat"),
                "up_cap":          extra.get("up_cap"),
                "dn_cap":          extra.get("dn_cap"),
                "n_alarms":        len(monitor.all_alarms),
                "alarm":           monitor.alarm,
                "alarm_date":      alarm_date,
                "alarm_date_ts":   alarm_date_ts,
                "all_alarm_dates": all_alarm_dates,
                "vix_regime":      vix_regime,
                "sig_quality":     sig_q,
                "dq_flags":        dq_flags,
                "L_now":           monitor.L,
                "h_threshold":     h_val,
            }

            if short_hist:
                insufficient_history.append({
                    "ticker": ticker,
                    "fund":   fund,
                    "bench":  bench,
                    "name":   name,
                    "months": n_months,
                })
                # Still add to main results so grid shows it (with note)

            all_rets.append(rets)
            all_monitors.append(mr)
            all_results.append(res)
            all_post_alarm.append(pa_df)

            tag = f"[{vix_regime}] " if vix_regime not in ("—","unknown") else ""
            print(f"OK  {tag}({mon_months}m live{'  SHORT' if short_hist else ''})")

        except Exception as e:
            print(f"ERROR: {e}")

    return all_rets, all_monitors, all_results, all_post_alarm, insufficient_history, vix_series


# ═══════════════════════════════════════════════════════════════════════════════
# CSV / JSON exports
# ═══════════════════════════════════════════════════════════════════════════════

def export_summary_csv(all_results, path):
    rows = []
    for r in all_results:
        rows.append({
            "Fund":       r["fund"],
            "Benchmark":  r["bench"],
            "Name":       r.get("name", ""),
            "AssetClass": _asset_class(r["fund"], r["bench"]),
            "Months":     r.get("months", ""),
            "IR":         r.get("ir", ""),
            "IR_lo":      r.get("ir_lo", ""),
            "IR_hi":      r.get("ir_hi", ""),
            "t_stat":     r.get("t_stat", ""),
            "TE":         r.get("te", ""),
            "HitRate":    r.get("hit_rate", ""),
            "UpCap":      r.get("up_cap", ""),
            "DnCap":      r.get("dn_cap", ""),
            "CumExcess":  r.get("cum_excess", ""),
            "L_now":      r.get("L_now", ""),
            "h":          r.get("h_threshold", ""),
            "N_alarms":   r.get("n_alarms", ""),
            "LastAlarm":  r.get("alarm_date", ""),
            "VIXRegime":  r.get("vix_regime", ""),
            "SigQuality": r.get("sig_quality", ""),
            "DQ_flags":   "; ".join(r.get("dq_flags", [])),
        })
    pd.DataFrame(rows).to_csv(path, index=False)
    print(f"  CSV saved: {path}")


def export_alarms_csv(all_results, all_post_alarm, path):
    rows = []
    for res, pa_df in zip(all_results, all_post_alarm):
        if not res["alarm"] or pa_df is None or pa_df.empty:
            continue
        for _, grp in pa_df.groupby("alarm_t"):
            alarm_t = grp["alarm_t"].iloc[0]
            alarm_dt = grp["alarm_date"].iloc[0]
            row12 = grp[grp["horizon"] == 12]
            row24 = grp[grp["horizon"] == 24]
            exc12 = row12["excess_cum"].iloc[0] if not row12.empty else float("nan")
            exc24 = row24["excess_cum"].iloc[0] if not row24.empty else float("nan")
            correct12 = row12["correct_alarm"].iloc[0] if not row12.empty else None
            verdict = ("pending" if correct12 is None or (isinstance(exc12, float) and np.isnan(exc12))
                       else ("correct" if correct12 else "false alarm"))
            rows.append({
                "Date":    alarm_dt,
                "Fund":    res["fund"],
                "Bench":   res["bench"],
                "h":       res.get("h_threshold", ""),
                "VIXRegime": res.get("vix_regime", ""),
                "+12M":    "" if np.isnan(exc12) else exc12,
                "+24M":    "" if np.isnan(exc24) else exc24,
                "Verdict": verdict,
            })
    pd.DataFrame(rows).sort_values("Date", ascending=False, na_position="last") \
      .to_csv(path, index=False)
    print(f"  Alarms CSV saved: {path}")


def export_health_json(all_results, all_post_alarm, path):
    n_funds  = len(all_results)
    n_alarms = sum(1 for r in all_results if r["alarm"])

    # Calm-regime signal quality
    sq_calm = None
    sq_vals = [r.get("sig_quality") for r in all_results
               if r.get("sig_quality") is not None and r.get("vix_regime") == "calm"]
    if sq_vals:
        sq_calm = float(np.mean(sq_vals))

    # Alarm counts by year
    alarm_by_year = {}
    for r in all_results:
        ts = r.get("alarm_date_ts")
        if ts is not None:
            yr = str(ts.year)
            alarm_by_year[yr] = alarm_by_year.get(yr, 0) + 1

    health = {
        "generated":          datetime.today().isoformat(),
        "n_funds":            n_funds,
        "n_alarms_total":     n_alarms,
        "calm_regime_sq_12m": sq_calm,
        "sq_target_pct":      0.60,
        "sq_below_target":    sq_calm is not None and sq_calm < 0.55,
        "alarm_by_year":      alarm_by_year,
        "expected_fa_per_yr": round(n_funds * 12 / 60, 2),
    }
    with open(path, "w") as f:
        json.dump(health, f, indent=2)
    print(f"  Health JSON saved: {path}")


# ═══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — low-level helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _hex_rgb(hex_str):
    """Convert a 6-digit hex colour string (no #) to RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _insert_field(run, field_code):
    """Inject a Word automatic field (PAGE or NUMPAGES) into an existing run."""
    fld_begin = _OxmlElement('w:fldChar')
    fld_begin.set(_qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = _OxmlElement('w:instrText')
    instr.set(_qn('xml:space'), 'preserve')
    instr.text = f' {field_code} '
    run._r.append(instr)
    fld_end = _OxmlElement('w:fldChar')
    fld_end.set(_qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


def _word_shade_cell(cell, hex_color):
    """Apply background fill colour to a single table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(_qn('w:shd')):
        tcPr.remove(existing)
    shd = _OxmlElement('w:shd')
    shd.set(_qn('w:val'),   'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'),  hex_color.lstrip('#').upper())
    tcPr.append(shd)


def _word_shade_row(row, hex_color):
    """Apply background fill to every cell in a table row."""
    for cell in row.cells:
        _word_shade_cell(cell, hex_color)


def _word_set_cell(cell, text, bold=False, color_hex=None,
                   size_pt=None, bg_hex=None, italic=False):
    """Set cell text with optional bold / colour / background formatting."""
    cell.text = ""
    p   = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name  = WORD_STYLE["body_font"]
    run.font.size  = Pt(size_pt or WORD_STYLE["table_pt"])
    run.bold       = bold
    run.italic     = italic
    if color_hex:
        run.font.color.rgb = _hex_rgb(color_hex)
    if bg_hex:
        _word_shade_cell(cell, bg_hex)


def _word_add_heading(doc, text, level=1):
    """Add a CUSUM-styled heading (navy Calibri, sized per WORD_STYLE)."""
    p  = doc.add_heading(text, level=level)
    pt = WORD_STYLE["heading1_pt"] if level == 1 else WORD_STYLE["heading2_pt"]
    for run in p.runs:
        run.font.name      = WORD_STYLE["heading_font"]
        run.font.size      = Pt(pt)
        run.font.color.rgb = _hex_rgb(WORD_STYLE["navy_hex"])
    return p


def _word_add_paragraph(doc, text, bold=False, italic=False,
                         color_hex=None, size_pt=None, style=None):
    """Add a body paragraph with optional formatting."""
    p   = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name  = WORD_STYLE["body_font"]
    run.font.size  = Pt(size_pt or WORD_STYLE["body_pt"])
    run.bold       = bold
    run.italic     = italic
    if color_hex:
        run.font.color.rgb = _hex_rgb(color_hex)
    return p


def _word_set_page_header(doc, data_cutoff):
    """Write 'CUSUM Active Manager Monitor | Data as of DATE' into every section header."""
    date_str    = (data_cutoff.strftime("%Y-%m-%d")
                   if hasattr(data_cutoff, 'strftime') else str(data_cutoff))
    header_text = WORD_STYLE["page_header_tpl"].format(date=date_str)
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        hdr_para = (section.header.paragraphs[0]
                    if section.header.paragraphs
                    else section.header.add_paragraph())
        hdr_para.text      = ""
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_para.add_run(header_text)
        run.font.name      = WORD_STYLE["header_font"]
        run.font.size      = Pt(8)
        run.font.color.rgb = _hex_rgb(WORD_STYLE["grey_hex"])


def _word_add_page_numbers(doc):
    """Add centred 'Page X of Y' footer to every section."""
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        ftr_para = (section.footer.paragraphs[0]
                    if section.footer.paragraphs
                    else section.footer.add_paragraph())
        ftr_para.text      = ""
        ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_pre = ftr_para.add_run("Page ")
        r_pre.font.size      = Pt(8)
        r_pre.font.color.rgb = _hex_rgb(WORD_STYLE["grey_hex"])

        r_pg = ftr_para.add_run()
        r_pg.font.size      = Pt(8)
        r_pg.font.color.rgb = _hex_rgb(WORD_STYLE["grey_hex"])
        _insert_field(r_pg, "PAGE")

        r_of = ftr_para.add_run(" of ")
        r_of.font.size      = Pt(8)
        r_of.font.color.rgb = _hex_rgb(WORD_STYLE["grey_hex"])

        r_tot = ftr_para.add_run()
        r_tot.font.size      = Pt(8)
        r_tot.font.color.rgb = _hex_rgb(WORD_STYLE["grey_hex"])
        _insert_field(r_tot, "NUMPAGES")


# ═══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — cover, executive summary, Part A
# ═══════════════════════════════════════════════════════════════════════════════

def _word_cover_page(doc, run_date, data_cutoff):
    """Cover page: title, subtitle, date metadata table, prepared-by line."""
    _word_add_paragraph(doc, "CUSUM Active Manager Monitor",
                        bold=True, size_pt=WORD_STYLE["cover_title_pt"],
                        color_hex=WORD_STYLE["navy_hex"])
    _word_add_paragraph(doc,
                        "Fund Universe Catalogue and Post-Alarm Performance Report",
                        size_pt=14, color_hex=WORD_STYLE["navy_hex"])
    doc.add_paragraph()

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    next_refresh = (pd.Timestamp(run_date) + pd.DateOffset(months=3)).strftime("%Y-%m-%d")
    rows_data = [
        ("Report date:",      run_date.strftime("%Y-%m-%d")),
        ("Data cutoff:",      data_cutoff.strftime("%Y-%m-%d")),
        ("Reporting period:", f"2005-01-01 – {data_cutoff.strftime('%Y-%m-%d')}"),
        ("Next refresh:",     next_refresh),
        ("Document version:", f"v{run_date.strftime('%Y-%m-%d')}"),
    ]
    for i, (label, value) in enumerate(rows_data):
        _word_set_cell(tbl.rows[i].cells[0], label, bold=True,
                       color_hex=WORD_STYLE["navy_hex"],
                       bg_hex=WORD_STYLE["lgrey_hex"])
        _word_set_cell(tbl.rows[i].cells[1], value)

    doc.add_paragraph()
    _word_add_paragraph(doc, "Prepared by: Investment Risk & Analytics",
                        italic=True, color_hex=WORD_STYLE["grey_hex"])
    _word_add_paragraph(doc, WORD_STYLE["confidential"],
                        italic=True, size_pt=8, color_hex=WORD_STYLE["grey_hex"])
    doc.add_page_break()


def _word_exec_summary(doc, all_results, all_post_alarm, vix_series):
    """Executive Summary: static narrative paragraphs + dynamic KPI paragraph."""
    _word_add_heading(doc, "Executive Summary", level=1)

    _word_add_paragraph(doc,
        "Purpose. This document gives senior management a complete view of the active "
        "fund universe under CUSUM monitoring, and the framework’s empirical track "
        "record. It is designed to support quarterly IC review and compliance sign-off.")
    _word_add_paragraph(doc,
        "Structure. The document is in two parts. Part A is the fund universe catalogue "
        "— every monitored mandate listed with ticker, asset class, benchmark, "
        "inception, AUM band, and typical tracking error. Part B is the post-alarm "
        "performance report: framework health metrics, the full alarm log, case study "
        "templates, and historical reference cases from the academic literature.")
    _word_add_paragraph(doc,
        "Why this matters. The framework’s credibility rests on its empirical track "
        "record, not just its theoretical pedigree. The alarm log and signal quality "
        "metrics in Part B are the evidence base for continued use of CUSUM as a "
        "governance tool.")
    _word_add_paragraph(doc,
        "Data note. Part A is fully populated and static. Part B figures (alarm counts, "
        "signal quality metrics, fund-level outcomes) are valid only as of the Data "
        "Cutoff date shown on the cover. Re-run report.py to refresh.")

    today    = datetime.today()
    cur_mo_s = pd.Timestamp(today.year, today.month, 1)
    cur_mo_e = cur_mo_s + pd.offsets.MonthEnd(0)

    in_alarm_now = [r for r in all_results
                    if r.get("L_now", 0) >= r.get("h_threshold", 19.46)]
    watchlist    = [r for r in all_results
                    if r.get("L_now", 0) >= 0.75 * r.get("h_threshold", 19.46)
                    and r not in in_alarm_now]
    new_this_mo  = [r for r in all_results
                    if r.get("alarm_date_ts") is not None
                    and cur_mo_s <= r["alarm_date_ts"] <= cur_mo_e]
    sq_vals  = [r["sig_quality"] for r in all_results
                if r.get("sig_quality") is not None
                and r.get("vix_regime") == "calm"]
    calm_sq  = f"{float(np.mean(sq_vals)):.0%}" if sq_vals else "n/a"

    kpi_text = (
        f"As of {today.strftime('%d %B %Y')}: {len(in_alarm_now)} fund(s) currently in "
        f"alarm state (L ≥ h), {len(watchlist)} on the watchlist "
        f"(L ≥ 0.75×h), and {len(new_this_mo)} new alarm(s) this month."
    )
    _word_add_paragraph(doc, kpi_text, bold=True,
                        color_hex=WORD_STYLE["navy_hex"])
    doc.add_page_break()


def _word_part_a_intro(doc):
    """Part A intro: purpose paragraph, 'How to read' bullets, composition note."""
    _word_add_heading(doc, "Part A — Fund Universe Catalogue", level=1)
    _word_add_paragraph(doc,
        "The framework monitors 67 active mandates across 13 asset class groupings. "
        "Each fund is paired with a peer-style benchmark index. The universe includes "
        "US large-cap, small-cap, developed international, balanced, EM equity, "
        "EM hard- and local-currency bond, global equity, and AIAIM Singapore-listed "
        "mandates.")

    _word_add_heading(doc, "How to read the table", level=2)
    _word_add_paragraph(doc, "Each fund row contains:")
    for bullet in [
        "Ticker — the symbol used in the production system (cusum.py PRESET_PAIRS)",
        "Fund Name — full legal name as reported in regulatory filings",
        "Asset Class — internal categorisation (drives sub-aggregation in monthly reports)",
        "Benchmark — the peer-style index used for excess return and tracking error",
        "Inception — first NAV available in our data history (approximate)",
        "AUM Band — Large (>$10B), Mid ($1–10B), Small (<$1B); informs investigation priority",
        "Typical TE — historical annualised tracking error range",
    ]:
        _word_add_paragraph(doc, bullet, style="List Bullet")

    _word_add_paragraph(doc,
        "Note on inception years and AUM bands. These are illustrative reference "
        "points based on publicly available fund information. Verify against official "
        "fund documents before use in formal reporting.")

    _word_add_heading(doc, "Universe composition", level=2)
    _word_add_paragraph(doc,
        "Concentration observation. The EM Equity sub-universe dominates the "
        "framework’s expected alarm rate. At ARL=60 months per fund, the universe "
        "generates approximately 13–14 expected false alarms per year in aggregate.")


_AC_NOTES: dict = {
    "US Large Blend":    "Most calibration-stable. Low serial correlation. SPY benchmark.",
    "US Large Growth":   "Russell 1000 Growth benchmark (IWF). Higher beta; TE 4–6%.",
    "US Large Value":    "Russell 1000 Value benchmark (IWD). Mean-reversion style.",
    "US Small Cap":      "Higher TE; noisier IR. Detection time ~21 months at h=19.46.",
    "US Small/Mid Cap":  "Broader mandate than pure small-cap. Fidelity FLPSX.",
    "Balanced 60/40":    "Lower TE → slower detection (~24 months). AOR benchmark.",
    "US Fixed Income":   "Lowest TE in universe. Numerically demanding. AGG benchmark.",
    "Intl Developed":    "Currency and benchmark choice matter. EFA (MSCI EAFE) benchmark.",
    "Intl Value":        "International value mandate; Dodge & Cox style.",
    "Emerging Markets":  "Largest sub-universe. Most alarms expected. EEM benchmark.",
    "EM Hard CCY Debt":  "Autocorrelation risk. EMB benchmark. em_debt CUSUM preset.",
    "EM Local CCY Debt": "Adds FX risk; widest TE bands. Small sample — use with caution.",
    "Global Equity":     "Includes ACWI and URTH benchmarks. Wide investment universe.",
}


def _word_part_a_composition_table(doc, all_results):
    """Asset class composition table: dynamic counts derived from PRESET_PAIRS."""
    ac_counts: dict = {}
    for ticker, (fund, bench, _name) in PRESET_PAIRS.items():
        ac = _asset_class(fund, bench)
        ac_counts[ac] = ac_counts.get(ac, 0) + 1

    total       = sum(ac_counts.values())
    ordered_acs = [ac for ac in ASSET_CLASS_ORDER if ac in ac_counts]
    for ac in ac_counts:
        if ac not in ordered_acs:
            ordered_acs.append(ac)

    tbl = doc.add_table(rows=1 + len(ordered_acs) + 1, cols=4)
    tbl.style = "Table Grid"

    _word_shade_row(tbl.rows[0], WORD_STYLE["navy_hex"])
    for i, h in enumerate(["Asset Class", "Count", "% of Universe", "Notes"]):
        _word_set_cell(tbl.rows[0].cells[i], h, bold=True,
                       color_hex="FFFFFF", size_pt=9)

    for ri, ac in enumerate(ordered_acs, start=1):
        count = ac_counts[ac]
        bg    = WORD_STYLE["lgrey_hex"] if ri % 2 == 0 else "FFFFFF"
        _word_set_cell(tbl.rows[ri].cells[0], ac,
                       bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[1], str(count),
                       bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[2], f"{count / total:.1%}",
                       bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[3], _AC_NOTES.get(ac, ""),
                       bg_hex=bg)

    last = len(ordered_acs) + 1
    _word_shade_row(tbl.rows[last], WORD_STYLE["lgrey_hex"])
    _word_set_cell(tbl.rows[last].cells[0], "TOTAL", bold=True)
    _word_set_cell(tbl.rows[last].cells[1], str(total), bold=True)
    _word_set_cell(tbl.rows[last].cells[2], "100.0%", bold=True)
    _word_set_cell(tbl.rows[last].cells[3],
                   f"Across {len(ordered_acs)} sub-categories")


def _word_part_a_universe_table(doc, all_results):
    """Full fund catalogue table: one row per fund, sorted by asset class then ticker."""
    _word_add_heading(doc, "Fund Universe — All 67 Mandates", level=1)
    _word_add_paragraph(doc,
        "Sorted by asset class then alphabetical ticker. "
        "AUM bands and typical TE ranges are illustrative; see FUND_METADATA in report.py.")

    rows_data = []
    for ticker, (fund, bench, name) in PRESET_PAIRS.items():
        ac              = _asset_class(fund, bench)
        meta            = FUND_METADATA.get(ticker, (None, "—", "—"))
        incept, aum, te = meta
        rows_data.append((_ac_sort_key(ac), ticker, name, ac, bench,
                          str(incept) if incept else "—", aum, te))
    rows_data.sort(key=lambda x: (x[0], x[1]))

    tbl = doc.add_table(rows=1 + len(rows_data), cols=7)
    tbl.style = "Table Grid"

    _word_shade_row(tbl.rows[0], WORD_STYLE["navy_hex"])
    for i, h in enumerate(["Ticker", "Fund Name", "Asset Class",
                            "Benchmark", "Incept.", "AUM", "Typ. TE"]):
        _word_set_cell(tbl.rows[0].cells[i], h, bold=True,
                       color_hex="FFFFFF", size_pt=9)

    for ri, (_sort, ticker, name, ac, bench, incept, aum, te) in \
            enumerate(rows_data, start=1):
        bg = WORD_STYLE["lgrey_hex"] if ri % 2 == 0 else "FFFFFF"
        _word_set_cell(tbl.rows[ri].cells[0], ticker, bold=True,
                       color_hex=WORD_STYLE["navy_hex"], bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[1], name,   bg_hex=bg, size_pt=8)
        _word_set_cell(tbl.rows[ri].cells[2], ac,     bg_hex=bg, size_pt=8)
        _word_set_cell(tbl.rows[ri].cells[3], bench,  bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[4], incept, bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[5], aum,    bg_hex=bg)
        _word_set_cell(tbl.rows[ri].cells[6], te,     bg_hex=bg)


# ═══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — Part B (health metrics, alarm log, case study template)
# ═══════════════════════════════════════════════════════════════════════════════

_TREND_FIRST_PERIOD = "First reporting period — trend data available from Q3 2026"


def _word_part_b1_health(doc, all_results, all_post_alarm):
    """B.1 Framework health metrics — count+ratio format (transparent-metric principle)."""
    _word_add_heading(doc, "Part B — Post-Alarm Performance Report", level=1)
    _word_add_paragraph(doc,
        "Purpose. This section is the framework's empirical track record. For each alarm "
        "fired by CUSUM, we document whether the fund subsequently underperformed its "
        "benchmark over the following 12 and 24 months. All percentages are shown with "
        "the underlying counts (transparent-metric principle, CUSUM_TODO.md).")

    _word_add_heading(doc, "B.1  Framework Health Metrics", level=2)
    _word_add_paragraph(doc,
        "Top-of-page summary each quarter. Signal quality pools all VIX regimes pending "
        "item c1 fix (CUSUM_TODO.md). Regime-stratified breakdown will replace these "
        "rows once c1 is resolved.")

    # ── Derive values ──────────────────────────────────────────────────────────
    yr_str    = str(datetime.today().year)
    ytd_total = sum(
        sum(1 for d in r.get("all_alarm_dates", []) if d.startswith(yr_str))
        for r in all_results
    )
    n_funds   = len(all_results)

    def _sq_counts(horizon):
        correct = wrong = pending = 0
        for pa_df in all_post_alarm:
            if pa_df is None or pa_df.empty:
                continue
            for _, row in pa_df[pa_df["horizon"] == horizon].iterrows():
                ca = row.get("correct_alarm")
                if ca is None or (isinstance(ca, float) and np.isnan(ca)):
                    pending += 1
                elif bool(ca):
                    correct += 1
                else:
                    wrong += 1
        return correct, wrong, pending

    c12, w12, p12 = _sq_counts(12)
    c24, w24, p24 = _sq_counts(24)
    c36, w36, p36 = _sq_counts(36)
    n12   = c12 + w12
    n24   = c24 + w24
    n36   = c36 + w36
    pct12 = f"{c12 / n12:.0%}" if n12 > 0 else "n/a"
    pct24 = f"{c24 / n24:.0%}" if n24 > 0 else "n/a"
    pct36 = f"{c36 / n36:.0%}" if n36 > 0 else "n/a"

    in_alarm    = sum(1 for r in all_results
                      if r.get("L_now", 0) >= r.get("h_threshold", 19.46))
    approaching = sum(1 for r in all_results
                      if (r.get("L_now", 0) >= 0.75 * r.get("h_threshold", 19.46)
                          and r.get("L_now", 0) < r.get("h_threshold", 19.46)))

    # ── Table (is_section_hdr, label, value, target) ──────────────────────────
    tbl_rows = [
        (True,  "YTD Alarm Activity",                                        "",       ""),
        (False, "Total alarms fired YTD",                                    str(ytd_total),
                f"~{round(n_funds * 12 / 60, 1)}/year expected"),
        (True,  "Signal Quality @ +12M  (all regimes — see note)",          "",       ""),
        (False, "Alarms with +12M outcome available",                        str(n12), "—"),
        (False, "  Of these: fund underperformed benchmark (correct signal)", str(c12), "—"),
        (False, "  Of these: fund outperformed benchmark (false alarm)",     str(w12), "—"),
        (False, "  Of these: +12M window not yet elapsed (pending)",         str(p12), "—"),
        (False, "Predictive accuracy @ +12M",                                pct12,   "≥ 60%"),
        (True,  "Signal Quality @ +24M  (all regimes — see note)",          "",       ""),
        (False, "Alarms with +24M outcome available",                        str(n24), "—"),
        (False, "  Of these: fund underperformed benchmark",                 str(c24), "—"),
        (False, "  Of these: fund outperformed benchmark",                   str(w24), "—"),
        (False, "Predictive accuracy @ +24M",                                pct24,   "≥ 55%"),
        (True,  "Signal Quality @ +36M  (all regimes — see note)",          "",       ""),
        (False, "Alarms with +36M outcome available",                        str(n36), "—"),
        (False, "  Of these: fund underperformed benchmark",                 str(c36), "—"),
        (False, "  Of these: fund outperformed benchmark",                   str(w36), "—"),
        (False, "Predictive accuracy @ +36M",                                pct36,   "≥ 50%"),
        (True,  "Current Universe Status",                                   "",       ""),
        (False, "Funds currently in alarm state  (L ≥ h)",                  str(in_alarm),    "≤ 3"),
        (False, "Funds approaching threshold  (L > 0.75 × h)",              str(approaching), "≤ 6"),
    ]

    tbl = doc.add_table(rows=1 + len(tbl_rows), cols=4)
    tbl.style = "Table Grid"

    _word_shade_row(tbl.rows[0], WORD_STYLE["navy_hex"])
    for ci, h in enumerate(["Metric", "Latest Value", "Target",
                             "Trend (last 4 quarters)"]):
        _word_set_cell(tbl.rows[0].cells[ci], h, bold=True,
                       color_hex="FFFFFF", size_pt=9)

    for ri, (is_hdr, label, value, target) in enumerate(tbl_rows, start=1):
        if is_hdr:
            _word_shade_row(tbl.rows[ri], WORD_STYLE["navy_hex"])
            _word_set_cell(tbl.rows[ri].cells[0], label, bold=True,
                           color_hex="FFFFFF", size_pt=8)
            for ci in range(1, 4):
                _word_set_cell(tbl.rows[ri].cells[ci], "",
                               bg_hex=WORD_STYLE["navy_hex"])
        else:
            bg     = WORD_STYLE["lgrey_hex"] if ri % 2 == 0 else "FFFFFF"
            is_pct = label.startswith("Predictive")
            _word_set_cell(tbl.rows[ri].cells[0], label,  bg_hex=bg, size_pt=9)
            _word_set_cell(tbl.rows[ri].cells[1], value,  bg_hex=bg, size_pt=9,
                           bold=is_pct,
                           color_hex=WORD_STYLE["navy_hex"] if is_pct else None)
            _word_set_cell(tbl.rows[ri].cells[2], target, bg_hex=bg, size_pt=9)
            _word_set_cell(tbl.rows[ri].cells[3], _TREND_FIRST_PERIOD,
                           bg_hex=bg, size_pt=8, italic=True,
                           color_hex=WORD_STYLE["grey_hex"])

    _word_add_paragraph(doc,
        "Note: Regime-stratified signal quality (calm / elevated / crisis breakdown) "
        "is pending item c1 fix (CUSUM_TODO.md). Current values pool all regimes. "
        "Remove this note once c1 is resolved.",
        italic=True, size_pt=8, color_hex=WORD_STYLE["amber_hex"])
    doc.add_page_break()


def _word_part_b2_alarms(doc, all_results, all_post_alarm, window_months=36):
    """B.2 Alarm log — every alarm event in the window, most recent first."""
    _word_add_heading(doc, "B.2  Alarm Log — Chronological", level=2)
    _word_add_paragraph(doc,
        f"Every CUSUM alarm in the most recent {window_months} months, most recent "
        "first. Compliance and IC can use this table to verify investigation status.")

    cutoff_ts = pd.Timestamp.today() - pd.DateOffset(months=window_months)
    events = []

    for res, pa_df in zip(all_results, all_post_alarm):
        if not res.get("alarm", False):
            continue
        h_val = res.get("h_threshold", 19.46)

        if pa_df is not None and not pa_df.empty:
            for _, grp in pa_df.groupby("alarm_t"):
                date_str = grp["alarm_date"].iloc[0]
                try:
                    alarm_ts = pd.Timestamp(date_str)
                except Exception:
                    continue
                if alarm_ts < cutoff_ts:
                    continue
                row12 = grp[grp["horizon"] == 12]
                row24 = grp[grp["horizon"] == 24]
                e12   = row12["excess_cum"].iloc[0] if not row12.empty else float("nan")
                e24   = row24["excess_cum"].iloc[0] if not row24.empty else float("nan")
                ca12  = row12["correct_alarm"].iloc[0] if not row12.empty else None
                v = ("pending"
                     if ca12 is None or (isinstance(e12, float) and np.isnan(e12))
                     else ("correct" if bool(ca12) else "false alarm"))
                events.append((alarm_ts, date_str, res["fund"], res["bench"],
                               h_val, res.get("vix_regime", "—"), e12, e24, v))
        else:
            alarm_ts = res.get("alarm_date_ts")
            if alarm_ts is None or alarm_ts < cutoff_ts:
                continue
            events.append((alarm_ts, alarm_ts.strftime("%Y-%m"),
                           res["fund"], res["bench"], h_val,
                           res.get("vix_regime", "—"),
                           float("nan"), float("nan"), "pending"))

    events.sort(key=lambda e: e[0], reverse=True)

    if not events:
        _word_add_paragraph(doc, "No alarms in the reporting window.", italic=True)
        return

    col_hdrs = ["Date", "Fund", "Bench", "h", "VIX Regime",
                "+12M Excess", "+24M Excess", "Verdict"]
    tbl = doc.add_table(rows=1 + len(events), cols=len(col_hdrs))
    tbl.style = "Table Grid"

    _word_shade_row(tbl.rows[0], WORD_STYLE["navy_hex"])
    for ci, h in enumerate(col_hdrs):
        _word_set_cell(tbl.rows[0].cells[ci], h, bold=True,
                       color_hex="FFFFFF", size_pt=8)

    for ri, (ts, date_str, fund, bench, h_val, regime, e12, e24, verdict) \
            in enumerate(events, start=1):
        bg    = WORD_STYLE["lgrey_hex"] if ri % 2 == 0 else "FFFFFF"
        e12_ok = isinstance(e12, float) and not np.isnan(e12)
        e24_ok = isinstance(e24, float) and not np.isnan(e24)
        e12_s  = f"{e12:+.1%}" if e12_ok else "pending"
        e24_s  = f"{e24:+.1%}" if e24_ok else "pending"
        v_col  = (WORD_STYLE["teal_hex"]  if verdict == "correct"    else
                  WORD_STYLE["coral_hex"] if verdict == "false alarm" else
                  WORD_STYLE["amber_hex"])
        r_col  = (WORD_STYLE["coral_hex"] if regime == "crisis"   else
                  WORD_STYLE["amber_hex"] if regime == "elevated" else
                  WORD_STYLE["teal_hex"])
        e12_col = (WORD_STYLE["teal_hex"]  if e12_ok and e12 < 0 else
                   WORD_STYLE["coral_hex"] if e12_ok and e12 > 0 else None)
        vals = [date_str, fund, bench, f"{h_val:.2f}",
                regime, e12_s, e24_s, verdict]
        clrs = [None, None, None, None, r_col, e12_col, None, v_col]
        for ci, (val, col) in enumerate(zip(vals, clrs)):
            _word_set_cell(tbl.rows[ri].cells[ci], val,
                           bg_hex=bg, color_hex=col, size_pt=8)
    doc.add_page_break()


def _word_part_b3_case_study(doc):
    """B.3 Case study template — static placeholder, one blank copy."""
    _word_add_heading(doc, "B.3  Case Study Template — Single Alarm Deep Dive", level=2)
    _word_add_paragraph(doc,
        "For each alarm with material portfolio implications, complete one copy of this "
        "template. Copy and duplicate as needed. Mechanical auto-population from "
        "cusum.py outputs is planned (CUSUM_TODO.md item d3).")

    for title, body in [
        ("Fund identification",
         "Ticker: [ticker]  |  Full fund name: [name]\n"
         "Benchmark: [ticker, full name]  |  Asset class: [category]\n"
         "AUM at alarm: [$X.X billion]"),
        ("Alarm event",
         "Alarm date: [YYYY-MM]  |  L at crossing: [19.xx]  |  h in force: [19.46 / 20.26 / 21.16]\n"
         "VIX regime: [Calm / Elevated / Crisis]"),
        ("Trailing context (12 months before alarm)",
         "Trailing IR: [+0.XX]  (95% CI [low, high])\n"
         "Trailing TE: [X.X% annualised]  |  Hit rate: [XX%]\n"
         "Up capture: [XXX%]   Down capture: [XXX%]"),
        ("Investigation findings",
         "Date of investigation: [YYYY-MM-DD]  |  Lead analyst: [name]\n"
         "  — PM continuity since last alarm: [yes / no / partial]\n"
         "  — AUM vs capacity: [within / approaching / above]\n"
         "  — Style consistency: [yes / drift detected]\n"
         "  — Operational red flags: [none / describe]"),
        ("Action taken",
         "Date of action: [YYYY-MM-DD or N/A]\n"
         "Action: [No action / Watch list / Reduced allocation / Redeemed]\n"
         "Rationale: [required for audit trail]"),
        ("Outcome tracking",
         "+6M:  Fund [X.X%]  Benchmark [Y.Y%]  Differential [Z.Z%]\n"
         "+12M post-alarm: [populate when window has elapsed]\n"
         "+24M post-alarm: [populate when window has elapsed]"),
        ("Lessons learned",
         "[Required before case study is marked complete. What did this alarm teach us "
         "about the framework, the fund, or our investigation process?]"),
    ]:
        _word_add_paragraph(doc, title, bold=True,
                            color_hex=WORD_STYLE["navy_hex"])
        _word_add_paragraph(doc, body, size_pt=9)
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT — Part B references, checklist, top-level orchestrator
# ═══════════════════════════════════════════════════════════════════════════════

def _word_part_b4_references(doc):
    """B.4 Historical reference cases — verbatim text from the original template."""
    _word_add_heading(doc, "B.4  Historical Reference Cases from the Literature",
                      level=2)
    _word_add_paragraph(doc,
        "Before our own track record is fully populated, the Morgan Stanley papers "
        "and the Philips-Yashchin-Stein 2003 paper provide validated historical cases "
        "that demonstrate CUSUM applied to known mutual fund decay episodes. These are "
        "referenced (not reproduced) below so management can see the framework has "
        "historical pedigree.")

    _word_add_heading(doc,
        "Case 1: Fidelity Magellan post-Lynch transition (referenced in PYS 2003)",
        level=2)
    _word_add_paragraph(doc,
        "Peter Lynch’s departure in 1990 marked the end of an extraordinary "
        "alpha-generating period for FMAGX. The PYS 2003 paper shows CUSUM detecting "
        "performance deterioration in the early-to-mid 1990s, several years before the "
        "underperformance became consensus view. Detection lag from regime change was "
        "within the framework’s predicted 21–32 month range. Reference: "
        "Philips, Yashchin & Stein (2003), Figure 3.")

    _word_add_heading(doc,
        "Case 2: Sequoia Fund (referenced in ‘When Good Managers Stumble’, MSIM)",
        level=2)
    _word_add_paragraph(doc,
        "SEQUX’s Valeant Pharmaceuticals concentration (2015–2016) led to "
        "severe underperformance. The MSIM 2018 paper discusses how CUSUM would have "
        "flagged this earlier than traditional methods, illustrating both the "
        "framework’s strengths and the importance of investigation context "
        "(Valeant was a single-name idiosyncratic event, not a slow skill decay). "
        "Reference: ‘When Good Managers Stumble’ (MSIM article).")

    _word_add_heading(doc,
        "Case 3: MSIM validation study (Display 3 of MSIM 2018)",
        level=2)
    _word_add_paragraph(doc,
        "The Morgan Stanley team validated CUSUM on the Morningstar US Large Cap, "
        "US Small Cap, and Global Large Cap universes from 1993–2020. Naive CUSUM "
        "achieved 52.8% predictive accuracy in identifying managers who would not "
        "recover in the three years following an alarm. With their additional "
        "peer-group filters layered on top, this rose to 59.5%. These published "
        "validation figures are useful benchmarks for what to expect from our own "
        "track record once it is populated. Reference: "
        "‘When Good Managers Stumble’, Display 4.")


def _word_part_b5_checklist(doc):
    """B.5 Quarterly review checklist — 7 bullet points, verbatim from template."""
    _word_add_heading(doc, "B.5  Quarterly Review Checklist", level=2)
    _word_add_paragraph(doc, "Each quarter, the Investment Risk team should:")

    for bullet in [
        "Re-run cusum.py with refreshed monthly NAV data through the latest "
        "complete month-end",
        "Populate Part B.1 (framework health metrics) with the latest values",
        "Append any new alarms to Part B.2 (alarm log) with current investigation "
        "status",
        "Complete a new case study (Part B.3 template) for each material alarm "
        "from the prior quarter",
        "Update outcome tracking for older alarms where new windows have elapsed "
        "(+12M, +24M, +36M)",
        "Review the signal quality trend; flag if calm-regime signal quality drops "
        "below 55%",
        "Re-confirm the threshold h is still appropriate; consider recalibration "
        "if h sensitivity flags drift",
    ]:
        _word_add_paragraph(doc, bullet, style="List Bullet")

    _word_add_paragraph(doc,
        "Document maintained by Investment Risk Team. Refreshed quarterly. "
        "All Part B numbers must be reproducible from cusum.py on the latest data "
        "— no exceptions.",
        italic=True, size_pt=9, color_hex=WORD_STYLE["grey_hex"])


def export_word_report(
    all_results,
    all_post_alarm,
    insufficient_history,
    vix_series,
    path,
    run_date=None,
    data_cutoff=None,
    window_months=36,
):
    """
    Generate the CUSUM management Word document from in-memory data.

    Signature per CONSOLIDATION_DESIGN.md. Do NOT call from main() until STEP 2.
    Returns the output path on success, None if python-docx is unavailable.
    """
    if not _DOCX_OK:
        print("  Word export skipped: python-docx not installed "
              "(pip install python-docx)")
        return None

    if run_date is None:
        run_date = datetime.today()
    if data_cutoff is None:
        data_cutoff = (
            pd.Timestamp.today() - 5 * pd.tseries.offsets.BusinessDay()
        ).to_pydatetime()

    doc = _DocxDocument()
    _word_set_page_header(doc, data_cutoff)
    _word_add_page_numbers(doc)

    _word_cover_page(doc, run_date, data_cutoff)
    _word_exec_summary(doc, all_results, all_post_alarm, vix_series)
    _word_part_a_intro(doc)
    _word_part_a_composition_table(doc, all_results)
    doc.add_page_break()
    _word_part_a_universe_table(doc, all_results)
    doc.add_page_break()
    _word_part_b1_health(doc, all_results, all_post_alarm)
    _word_part_b2_alarms(doc, all_results, all_post_alarm,
                         window_months=window_months)
    _word_part_b3_case_study(doc)
    doc.add_page_break()
    _word_part_b4_references(doc)
    _word_part_b5_checklist(doc)

    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)
    print(f"  Word report saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="CUSUM Active Manager Monitor — Report Generator")
    parser.add_argument("--output",       default=None,
                        help="Output PDF path (default: output/cusum_report.pdf)")
    parser.add_argument("--refresh-all",  action="store_true",
                        help="Re-download all data and refresh cache")
    parser.add_argument("--use-cache",    action="store_true",
                        help="Load from cache if available, else fetch")
    parser.add_argument("--recalibrate",  action="store_true",
                        help="Recalibrate thresholds before running (slow)")
    parser.add_argument("--skip-word",    action="store_true",
                        help="Skip Word document export (PDF only)")
    args = parser.parse_args()

    out_dir  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(out_dir, exist_ok=True)
    pdf_path   = args.output or os.path.join(out_dir, "cusum_report.pdf")
    cache_path = os.path.join(out_dir, "cusum_data_cache.pkl")

    if args.recalibrate:
        from cusum import calibrate_all_presets
        print("Recalibrating thresholds (Monte Carlo)...")
        calibrate_all_presets()

    # ── Load or collect data ─────────────────────────────────────────────────
    loaded_from_cache = False
    if args.use_cache and not args.refresh_all and os.path.exists(cache_path):
        try:
            with open(cache_path, "rb") as f:
                cached = pickle.load(f)
            all_rets, all_monitors, all_results, all_post_alarm, \
                insufficient_history, vix_series = cached
            loaded_from_cache = True
            print(f"Loaded from cache: {cache_path}  ({len(all_results)} funds)")
        except Exception as e:
            print(f"Cache load failed ({e}), re-fetching...")

    if not loaded_from_cache:
        (all_rets, all_monitors, all_results, all_post_alarm,
         insufficient_history, vix_series) = collect_data()

        if args.use_cache or args.refresh_all:
            try:
                with open(cache_path, "wb") as f:
                    pickle.dump((all_rets, all_monitors, all_results, all_post_alarm,
                                 insufficient_history, vix_series), f)
                print(f"  Cache saved: {cache_path}")
            except Exception as e:
                print(f"  Cache save failed: {e}")

    if not all_results:
        print("No data collected — aborting.")
        return

    total_pages = compute_total_pages(all_results, insufficient_history)
    pnum        = [1]

    print(f"\nBuilding PDF -> {pdf_path}  ({total_pages} pages expected)")
    with PdfPages(pdf_path) as pdf:
        page_cover             (pdf, pnum, total_pages)
        page_executive_summary (pdf, all_results, all_post_alarm, vix_series, pnum, total_pages)
        page_methodology       (pdf, pnum, total_pages)
        page_assumptions       (pdf, pnum, total_pages)
        page_fund_grid         (pdf, all_rets, all_monitors, all_results, pnum, total_pages)
        page_summary_table     (pdf, all_results, pnum, total_pages)
        page_alarm_log         (pdf, all_results, all_post_alarm, pnum, total_pages)
        page_framework_health  (pdf, all_results, all_post_alarm, pnum, total_pages)
        page_waterfall         (pdf, all_results, pnum, total_pages)
        page_insufficient_history(pdf, insufficient_history, pnum, total_pages)
        page_data_quality      (pdf, all_results, pnum, total_pages)
        page_conclusion        (pdf, all_results, all_post_alarm, pnum, total_pages)

    print(f"  PDF saved: {pdf_path}  ({pnum[0]-1} pages written)")

    # ── Export CSV/JSON to dated run directory ────────────────────────────────
    run_dir   = os.path.join(out_dir, "cusum_runs",
                             datetime.today().strftime("%Y-%m-%d"))
    os.makedirs(run_dir, exist_ok=True)
    csv_path  = os.path.join(run_dir, "fund_results.csv")
    alm_path  = os.path.join(run_dir, "alarm_log.csv")
    hlth_path = os.path.join(run_dir, "health.json")

    try:
        export_summary_csv(all_results, csv_path)
    except Exception as e:
        print(f"  CSV export failed: {e}")
    try:
        export_alarms_csv(all_results, all_post_alarm, alm_path)
    except Exception as e:
        print(f"  Alarms CSV failed: {e}")
    try:
        export_health_json(all_results, all_post_alarm, hlth_path)
    except Exception as e:
        print(f"  Health JSON failed: {e}")

    # ── Word document export ──────────────────────────────────────────────────
    if not args.skip_word:
        today_str = datetime.today().strftime("%Y%m%d")
        reports_dir = os.path.join(out_dir, "reports")
        os.makedirs(reports_dir, exist_ok=True)
        word_path = os.path.join(reports_dir, f"CUSUM_Report_{today_str}.docx")
        try:
            export_word_report(
                all_results, all_post_alarm, insufficient_history, vix_series,
                path=word_path,
            )
        except Exception as e:
            print(f"  Word export failed: {e}")

    print("Done.")


if __name__ == "__main__":
    main()
